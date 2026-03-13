#!/usr/bin/env python3
"""Generate the MCP PCB EMCopilot Framework reference document (DOCX).

Creates a comprehensive engineering reference covering architecture,
tool categories, design review workflow, and API reference.

Generates diagrams as SVG using Python, converts to PNG via cairosvg,
and embeds them in the DOCX using python-docx.
"""

import io
import os
import math
import textwrap
from datetime import date

import cairosvg
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "MCP_PCB_EMCopilot_Framework.docx")
DIAGRAM_DIR = os.path.join(SCRIPT_DIR, "diagrams")

DOCUMENT_REV = "1.0"
DOCUMENT_DATE = date.today().isoformat()

# ---------------------------------------------------------------------------
# SVG Diagram Generation
# ---------------------------------------------------------------------------

def _svg_header(width, height):
    return (f'<svg xmlns="http://www.w3.org/2000/svg" '
            f'width="{width}" height="{height}" '
            f'viewBox="0 0 {width} {height}">\n'
            f'<style>\n'
            f'  text {{ font-family: Arial, Helvetica, sans-serif; }}\n'
            f'  .title {{ font-size: 18px; font-weight: bold; fill: #1a1a1a; }}\n'
            f'  .subtitle {{ font-size: 11px; fill: #666; }}\n'
            f'  .label {{ font-size: 11px; font-weight: bold; fill: #fff; }}\n'
            f'  .label-dark {{ font-size: 11px; font-weight: bold; fill: #333; }}\n'
            f'  .label-sm {{ font-size: 9px; fill: #fff; }}\n'
            f'  .label-sm-dark {{ font-size: 9px; fill: #555; }}\n'
            f'  .dim {{ font-size: 9px; fill: #888; }}\n'
            f'  .arrow {{ stroke: #444; stroke-width: 2; fill: none; marker-end: url(#arrowhead); }}\n'
            f'  .arrow-dash {{ stroke: #999; stroke-width: 1.5; stroke-dasharray: 6,3; fill: none; marker-end: url(#arrowhead-gray); }}\n'
            f'</style>\n'
            f'<defs>\n'
            f'  <marker id="arrowhead" markerWidth="10" markerHeight="7" refX="10" refY="3.5" orient="auto">\n'
            f'    <polygon points="0 0, 10 3.5, 0 7" fill="#444" />\n'
            f'  </marker>\n'
            f'  <marker id="arrowhead-gray" markerWidth="8" markerHeight="6" refX="8" refY="3" orient="auto">\n'
            f'    <polygon points="0 0, 8 3, 0 6" fill="#999" />\n'
            f'  </marker>\n'
            f'</defs>\n')


def _rounded_rect(x, y, w, h, fill, stroke, rx=6):
    return (f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{rx}" '
            f'fill="{fill}" stroke="{stroke}" stroke-width="1.5"/>\n')


def _escape_xml(s):
    """Escape XML special characters."""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _text(x, y, text, css_class="label", anchor="middle"):
    return f'<text x="{x}" y="{y}" class="{css_class}" text-anchor="{anchor}">{_escape_xml(text)}</text>\n'


def _line(x1, y1, x2, y2, css_class="arrow"):
    return f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" class="{css_class}"/>\n'


def generate_architecture_svg():
    """Generate the main architecture diagram."""
    W, H = 1000, 720
    svg = _svg_header(W, H)

    # Title
    svg += _text(W // 2, 30, "MCP PCB EMCopilot - System Architecture", "title")

    # --- Row 1: MCP Clients ---
    clients = ["Claude Code", "Codex CLI", "Other MCP\nClients"]
    client_colors = [("#2563eb", "#1d4ed8"), ("#7c3aed", "#6d28d9"), ("#64748b", "#475569")]
    cx_start = 300
    cx_spacing = 160
    cy = 70
    cw, ch = 130, 44

    # Group label
    svg += _rounded_rect(cx_start - 20, cy - 10, len(clients) * cx_spacing + 10, ch + 30,
                         "#f0f4ff", "#93c5fd", rx=10)
    svg += _text(cx_start - 5, cy + ch + 15, "MCP Clients", "label-sm-dark", "start")

    for i, (name, (fill, stroke)) in enumerate(zip(clients, client_colors)):
        x = cx_start + i * cx_spacing
        svg += _rounded_rect(x, cy, cw, ch, fill, stroke)
        lines = name.split("\n")
        for j, line in enumerate(lines):
            svg += _text(x + cw // 2, cy + 20 + j * 14, line, "label")

    # --- Row 2: MCP Server ---
    sy = 175
    sw, sh = 350, 50
    sx = W // 2 - sw // 2
    svg += _rounded_rect(sx, sy, sw, sh, "#0f766e", "#0d9488", rx=8)
    svg += _text(W // 2, sy + 22, "MCP Server (93 Tools)", "label")
    svg += _text(W // 2, sy + 38, "Session Manager / Protocol Layer", "label-sm")

    # Arrows from clients to server
    for i in range(3):
        cx = cx_start + i * cx_spacing + cw // 2
        svg += _line(cx, cy + ch, W // 2, sy)

    # --- Row 3: Parser Layer ---
    py = 290
    parsers = [
        ("KiCad", ".kicad_pcb"),
        ("ODB++", ".zip/.tgz"),
        ("Gerber", "RS-274X"),
        ("IPC-2581", ".xml"),
        ("Altium", ".PcbDoc"),
        ("STEP", ".step/.stp"),
        ("PDF Schem.", "PDF"),
        ("BOM", ".csv/.xlsx"),
    ]
    pw, ph = 100, 48
    px_start = (W - len(parsers) * (pw + 12)) // 2 + 6

    # Parser group box
    svg += _rounded_rect(px_start - 15, py - 15, len(parsers) * (pw + 12) + 18, ph + 45,
                         "#fef3c7", "#fbbf24", rx=10)
    svg += _text(px_start, py + ph + 25, "Parser Layer", "label-sm-dark", "start")

    for i, (name, fmt) in enumerate(parsers):
        x = px_start + i * (pw + 12)
        svg += _rounded_rect(x, py, pw, ph, "#f59e0b", "#d97706", rx=4)
        svg += _text(x + pw // 2, py + 18, name, "label")
        svg += _text(x + pw // 2, py + 34, fmt, "label-sm")

    # Arrow from server to parser layer
    svg += _line(W // 2, sy + sh, W // 2, py - 15)

    # --- Row 4: Analyzer Layer ---
    ay = 415
    analyzers = [
        ("EMC / EMI", "#dc2626", "#b91c1c"),
        ("Signal\nIntegrity", "#2563eb", "#1d4ed8"),
        ("Power\nIntegrity", "#ea580c", "#c2410c"),
        ("Thermal", "#ca8a04", "#a16207"),
        ("DFM", "#16a34a", "#15803d"),
        ("RF/Antenna", "#7c3aed", "#6d28d9"),
        ("ESD", "#db2777", "#be185d"),
        ("High-Speed\nDigital", "#0891b2", "#0e7490"),
    ]
    aw, ah = 100, 52
    ax_start = (W - len(analyzers) * (aw + 12)) // 2 + 6

    # Analyzer group box
    svg += _rounded_rect(ax_start - 15, ay - 15, len(analyzers) * (aw + 12) + 18, ah + 45,
                         "#fce7f3", "#f472b6", rx=10)
    svg += _text(ax_start, ay + ah + 25, "Analyzer Layer", "label-sm-dark", "start")

    for i, (name, fill, stroke) in enumerate(analyzers):
        x = ax_start + i * (aw + 12)
        svg += _rounded_rect(x, ay, aw, ah, fill, stroke, rx=4)
        lines = name.split("\n")
        for j, line in enumerate(lines):
            svg += _text(x + aw // 2, ay + 22 + j * 14, line, "label")

    # Arrows from parser to analyzer
    svg += _line(W // 2, py + ph + 30, W // 2, ay - 15)

    # --- Row 5: Output Layer ---
    oy = 545
    outputs = [
        ("SVG Renders", "#0d9488", "#0f766e"),
        ("PNG Export", "#3b82f6", "#2563eb"),
        ("DOCX Reports", "#8b5cf6", "#7c3aed"),
        ("Annotated\nBoards", "#f97316", "#ea580c"),
    ]
    ow, oh = 130, 48
    ox_start = (W - len(outputs) * (ow + 20)) // 2 + 10

    # Output group box
    svg += _rounded_rect(ox_start - 15, oy - 15, len(outputs) * (ow + 20) + 10, oh + 45,
                         "#ecfdf5", "#6ee7b7", rx=10)
    svg += _text(ox_start, oy + oh + 25, "Report & Visualization Layer", "label-sm-dark", "start")

    for i, (name, fill, stroke) in enumerate(outputs):
        x = ox_start + i * (ow + 20)
        svg += _rounded_rect(x, oy, ow, oh, fill, stroke, rx=4)
        lines = name.split("\n")
        for j, line in enumerate(lines):
            svg += _text(x + ow // 2, oy + 22 + j * 14, line, "label")

    # Arrow from analyzer to output
    svg += _line(W // 2, ay + ah + 30, W // 2, oy - 15)

    # --- Integration sidebar ---
    ix, iy = 830, 290
    iw, ih_item = 155, 36
    integrations = [
        ("mcp-emc-regulations", "#6366f1"),
        ("mcp-nec2-antenna", "#8b5cf6"),
        ("mcp-openems", "#a855f7"),
        ("mcp-drawio-eng.", "#c084fc"),
    ]
    svg += _rounded_rect(ix - 10, iy - 25, iw + 20, len(integrations) * (ih_item + 8) + 40,
                         "#f5f3ff", "#c4b5fd", rx=8)
    svg += _text(ix + iw // 2, iy - 5, "Integrations", "label-dark")
    for i, (name, fill) in enumerate(integrations):
        y = iy + 15 + i * (ih_item + 8)
        svg += _rounded_rect(ix, y, iw, ih_item, fill, fill, rx=4)
        svg += _text(ix + iw // 2, y + 22, name, "label")
    # Arrow from integrations to server
    svg += f'<line x1="{ix}" y1="{iy + 80}" x2="{sx + sw}" y2="{sy + sh // 2}" class="arrow-dash"/>\n'

    svg += '</svg>\n'
    return svg


def generate_stackup_svg():
    """Generate a 6-layer PCB stackup cross-section diagram."""
    W, H = 700, 420
    svg = _svg_header(W, H)
    svg += _text(W // 2, 28, "6-Layer PCB Stackup Cross-Section", "title")

    layers = [
        ("Soldermask", 6, "#009900", "#fff", "0.025 mm"),
        ("L1 Signal (1oz Cu)", 10, "#cc3333", "#fff", "0.035 mm"),
        ("Prepreg (FR-4)", 25, "#ffffcc", "#666", "0.150 mm, Dk=4.2"),
        ("L2 GND (1oz Cu)", 10, "#3333cc", "#fff", "0.035 mm"),
        ("Core (FR-4)", 45, "#e0e0e0", "#666", "0.300 mm, Dk=4.2"),
        ("L3 Signal (0.5oz Cu)", 8, "#cc3333", "#fff", "0.018 mm"),
        ("Prepreg (FR-4)", 45, "#ffffcc", "#666", "0.300 mm, Dk=4.2"),
        ("L4 Signal (0.5oz Cu)", 8, "#cc3333", "#fff", "0.018 mm"),
        ("Core (FR-4)", 45, "#e0e0e0", "#666", "0.300 mm, Dk=4.2"),
        ("L5 PWR (1oz Cu)", 10, "#cc6600", "#fff", "0.035 mm"),
        ("Prepreg (FR-4)", 25, "#ffffcc", "#666", "0.150 mm, Dk=4.2"),
        ("L6 Signal (1oz Cu)", 10, "#cc3333", "#fff", "0.035 mm"),
        ("Soldermask", 6, "#009900", "#fff", "0.025 mm"),
    ]

    lx, lw = 150, 350
    ly = 55
    for name, h, fill, text_color, dims in layers:
        svg += (f'<rect x="{lx}" y="{ly}" width="{lw}" height="{h}" '
                f'fill="{fill}" stroke="#888" stroke-width="0.5"/>\n')
        # Label on left
        label_class = "label-sm-dark" if text_color == "#666" else "label-sm-dark"
        svg += _text(lx - 10, ly + h // 2 + 4, name, "label-sm-dark", "end")
        # Dimensions on right
        svg += _text(lx + lw + 10, ly + h // 2 + 4, dims, "dim", "start")
        ly += h

    # Total thickness bracket
    bracket_x = lx + lw + 130
    svg += f'<line x1="{bracket_x}" y1="55" x2="{bracket_x}" y2="{ly}" stroke="#cc0000" stroke-width="2"/>\n'
    svg += f'<line x1="{bracket_x - 5}" y1="55" x2="{bracket_x + 5}" y2="55" stroke="#cc0000" stroke-width="2"/>\n'
    svg += f'<line x1="{bracket_x - 5}" y1="{ly}" x2="{bracket_x + 5}" y2="{ly}" stroke="#cc0000" stroke-width="2"/>\n'
    svg += f'<text x="{bracket_x + 10}" y="{(55 + ly) // 2 + 5}" font-size="12" font-weight="bold" fill="#cc0000" font-family="Arial">Total: ~1.43 mm</text>\n'

    svg += '</svg>\n'
    return svg


def generate_rf_analysis_svg():
    """Generate RF/EMI analysis flow diagram."""
    W, H = 800, 350
    svg = _svg_header(W, H)
    svg += _text(W // 2, 28, "EMI Analysis Signal Flow", "title")

    blocks = [
        ("Clock/Signal\nSource", 50, 120, 110, 55, "#7c3aed", "#6d28d9"),
        ("PCB Trace\n(Transmission Line)", 200, 120, 120, 55, "#2563eb", "#1d4ed8"),
        ("Unintentional\nAntenna", 360, 120, 110, 55, "#dc2626", "#b91c1c"),
        ("Radiated\nEmission", 510, 120, 110, 55, "#ea580c", "#c2410c"),
        ("Regulatory\nLimit Check", 660, 120, 110, 55, "#16a34a", "#15803d"),
    ]

    for name, x, y, w, h, fill, stroke in blocks:
        svg += _rounded_rect(x, y, w, h, fill, stroke)
        lines = name.split("\n")
        for j, line in enumerate(lines):
            svg += _text(x + w // 2, y + 22 + j * 14, line, "label")

    # Arrows between blocks
    arrows = [(160, 147), (320, 147), (470, 147), (620, 147)]
    for ax, ay in arrows:
        svg += _line(ax, ay, ax + 40, ay)

    # Annotation boxes below
    annotations = [
        ("Harmonic\nContent", 65, 220, 90, 40, "#f5f3ff", "#7c3aed"),
        ("Impedance &\nLoss Model", 210, 220, 100, 40, "#eff6ff", "#2563eb"),
        ("Loop Area &\nTrace Length", 370, 220, 95, 40, "#fef2f2", "#dc2626"),
        ("E-field\n(dBuV/m)", 520, 220, 90, 40, "#fff7ed", "#ea580c"),
        ("FCC/CISPR\nMargin (dB)", 670, 220, 95, 40, "#f0fdf4", "#16a34a"),
    ]
    for name, x, y, w, h, fill, stroke in annotations:
        svg += _rounded_rect(x, y, w, h, fill, stroke, rx=4)
        lines = name.split("\n")
        for j, line in enumerate(lines):
            svg += _text(x + w // 2, y + 16 + j * 13, line, "label-sm-dark")

    # Dashed arrows from annotations to blocks
    for i, (_, ax, ay, aw, ah, _, _) in enumerate(annotations):
        bx = blocks[i][1] + blocks[i][3] // 2
        by = blocks[i][2] + blocks[i][4]
        svg += f'<line x1="{ax + aw // 2}" y1="{ay}" x2="{bx}" y2="{by}" class="arrow-dash"/>\n'

    # Formula note
    svg += _text(W // 2, 310, "Based on: Trapezoidal waveform Fourier analysis, IPC-2141 impedance, magnetic dipole radiation model", "dim")

    svg += '</svg>\n'
    return svg


def generate_emc_test_svg():
    """Generate an EMC test setup diagram (CISPR 25 RE)."""
    W, H = 800, 500
    svg = _svg_header(W, H)
    svg += _text(W // 2, 28, "CISPR 25 Radiated Emissions Test Setup", "title")
    svg += _text(W // 2, 48, "Semi-anechoic chamber | Antenna 1m from EUT | 150 kHz - 2.5 GHz", "dim")

    # Ground plane
    svg += f'<rect x="50" y="400" width="700" height="15" fill="#ccc" stroke="#999" stroke-width="1"/>\n'
    svg += _text(400, 412, "Ground Plane", "label-sm-dark")

    # Test table
    svg += f'<rect x="120" y="250" width="160" height="150" fill="none" stroke="#666" stroke-width="2" rx="2"/>\n'
    svg += _text(200, 395, "Test Table (0.8m)", "dim")

    # EUT
    svg += _rounded_rect(145, 200, 110, 55, "#dbeafe", "#3b82f6", rx=4)
    svg += _text(200, 232, "EUT (PCB)", "label-dark")

    # Antenna mast
    svg += f'<rect x="415" y="180" width="6" height="220" fill="#888" stroke="#666"/>\n'

    # Antenna
    svg += f'<polygon points="418,180 400,220 436,220" fill="#f5f5f5" stroke="#666" stroke-width="1.5"/>\n'
    svg += _text(418, 175, "ANT", "label-dark")

    # LISN
    svg += _rounded_rect(60, 310, 90, 40, "#fed7aa", "#ea580c", rx=6)
    svg += _text(105, 335, "LISN", "label-dark")

    # Spectrum Analyzer
    svg += _rounded_rect(560, 100, 120, 50, "#e9d5ff", "#7c3aed", rx=6)
    svg += _text(620, 122, "Spectrum", "label-dark")
    svg += _text(620, 136, "Analyzer", "label-dark")

    # Cables
    svg += f'<line x1="150" y1="330" x2="145" y2="255" stroke="#333" stroke-width="1.5"/>\n'
    svg += f'<line x1="560" y1="125" x2="436" y2="200" stroke="#333" stroke-width="1.5"/>\n'
    svg += f'<line x1="150" y1="330" x2="560" y2="125" stroke="#999" stroke-width="1" stroke-dasharray="5,3"/>\n'

    # Distance annotation
    svg += f'<line x1="200" y1="440" x2="418" y2="440" stroke="#cc0000" stroke-width="1" stroke-dasharray="4,3"/>\n'
    svg += f'<polygon points="200,437 200,443 208,440" fill="#cc0000"/>\n'
    svg += f'<polygon points="418,437 418,443 410,440" fill="#cc0000"/>\n'
    svg += f'<rect x="280" y="432" width="50" height="16" fill="white"/>\n'
    svg += _text(305, 445, "1 m", "label-sm-dark")

    # Absorber wedges (simplified)
    for y in range(70, 400, 40):
        svg += f'<polygon points="30,{y} 45,{y+15} 30,{y+30}" fill="#555" opacity="0.3"/>\n'
        svg += f'<polygon points="770,{y} 755,{y+15} 770,{y+30}" fill="#555" opacity="0.3"/>\n'

    svg += '</svg>\n'
    return svg


def generate_design_review_workflow_svg():
    """Generate a design review workflow diagram."""
    W, H = 900, 280
    svg = _svg_header(W, H)
    svg += _text(W // 2, 28, "Automated Design Review Workflow", "title")

    steps = [
        ("1. Parse\nLayout", "#f59e0b"),
        ("2. Classify\nDesign", "#3b82f6"),
        ("3. Detect\nInterfaces", "#8b5cf6"),
        ("4. Run\nAnalyzers", "#dc2626"),
        ("5. Cross-\nCorrelate", "#ea580c"),
        ("6. Generate\nVisualizations", "#0d9488"),
        ("7. Export\nReport", "#16a34a"),
    ]

    sw, sh = 105, 55
    sx_start = 30
    sx_gap = 125

    for i, (name, fill) in enumerate(steps):
        x = sx_start + i * sx_gap
        y = 70
        svg += _rounded_rect(x, y, sw, sh, fill, fill, rx=6)
        lines = name.split("\n")
        for j, line in enumerate(lines):
            svg += _text(x + sw // 2, y + 22 + j * 14, line, "label")
        if i < len(steps) - 1:
            svg += _line(x + sw, y + sh // 2, x + sx_gap, y + sh // 2)

    # Sub-labels
    sublabels = [
        "KiCad / ODB++\nGerber / Altium",
        "RF / Mixed-Sig\nHigh-Speed / Power",
        "DDR / PCIe / USB\nEthernet / LVDS",
        "EMC / SI / PI\nDFM / Thermal",
        "Related findings\nSeverity ranking",
        "Board renders\nNet highlights",
        "DOCX with\nembedded images",
    ]
    for i, sublabel in enumerate(sublabels):
        x = sx_start + i * sx_gap + sw // 2
        lines = sublabel.split("\n")
        for j, line in enumerate(lines):
            svg += _text(x, 155 + j * 14, line, "dim")

    svg += '</svg>\n'
    return svg


def svg_to_png_bytes(svg_string, scale=2):
    """Convert SVG string to PNG bytes using cairosvg."""
    return cairosvg.svg2png(bytestring=svg_string.encode('utf-8'), scale=scale)


def generate_all_diagrams():
    """Generate all diagram PNGs. Returns dict of name -> PNG bytes."""
    diagrams = {}
    generators = {
        "architecture": generate_architecture_svg,
        "stackup": generate_stackup_svg,
        "rf_analysis": generate_rf_analysis_svg,
        "emc_test": generate_emc_test_svg,
        "workflow": generate_design_review_workflow_svg,
    }
    os.makedirs(DIAGRAM_DIR, exist_ok=True)
    for name, gen_func in generators.items():
        svg = gen_func()
        # Save SVG for reference
        svg_path = os.path.join(DIAGRAM_DIR, f"{name}.svg")
        with open(svg_path, 'w') as f:
            f.write(svg)
        # Convert to PNG
        png_bytes = svg_to_png_bytes(svg)
        png_path = os.path.join(DIAGRAM_DIR, f"{name}.png")
        with open(png_path, 'wb') as f:
            f.write(png_bytes)
        diagrams[name] = png_bytes
        print(f"  Generated {name}.png ({len(png_bytes)} bytes)")
    return diagrams


# ---------------------------------------------------------------------------
# DOCX Styling Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FA")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")
    return table


def add_image_from_bytes(doc, png_bytes, caption, width_inches=6.0):
    """Add a PNG image from bytes with a caption."""
    stream = io.BytesIO(png_bytes)
    doc.add_picture(stream, width=Inches(width_inches))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# Document Generation
# ---------------------------------------------------------------------------

def create_cover_page(doc):
    """Create the cover page."""
    # Add blank paragraphs for spacing
    for _ in range(6):
        doc.add_paragraph("")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MCP PCB EMCopilot Framework")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Comprehensive PCB Design Review,\nEMC Analysis & Signal Integrity Platform")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    # Horizontal rule
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph("")

    # Metadata table
    meta = [
        ("Version", f"v0.2.0 (Document Rev {DOCUMENT_REV})"),
        ("Date", DOCUMENT_DATE),
        ("Author", "Adam Engelbrecht (RFingAdam)"),
        ("License", "Apache-2.0"),
        ("Repository", "github.com/RFingAdam/mcp-pcb-emcopilot"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, val) in enumerate(meta):
        table.rows[i].cells[0].text = ""
        table.rows[i].cells[1].text = ""
        p0 = table.rows[i].cells[0].paragraphs[0]
        r0 = p0.add_run(key)
        r0.bold = True
        r0.font.size = Pt(11)
        r0.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p1 = table.rows[i].cells[1].paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(11)

    # Remove table borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="none"/>'
                '  <w:left w:val="none"/>'
                '  <w:bottom w:val="none"/>'
                '  <w:right w:val="none"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)

    doc.add_page_break()


def add_toc(doc):
    """Add a Table of Contents field."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TABLE OF CONTENTS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph("")

    # Add TOC field code
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run("Right-click and select 'Update Field' to populate Table of Contents")
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run4.font.size = Pt(10)
    run4.italic = True

    run5 = paragraph.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    doc.add_page_break()


def add_section_heading(doc, text, level=1):
    """Add a heading with consistent styling."""
    doc.add_heading(text, level=level)


def add_body_text(doc, text):
    """Add body text paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p


def add_code_block(doc, code, language=""):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Light gray background via shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)


# ---------------------------------------------------------------------------
# Section: Framework Overview
# ---------------------------------------------------------------------------

def section_overview(doc):
    add_section_heading(doc, "1. Framework Overview")

    add_section_heading(doc, "1.1 What is MCP PCB EMCopilot", level=2)
    add_body_text(doc,
        "MCP PCB EMCopilot is a Model Context Protocol (MCP) server that provides 93 specialized tools "
        "for AI-assisted PCB design review, electromagnetic compatibility (EMC) analysis, signal integrity (SI) "
        "verification, and power integrity (PI) assessment. It enables AI assistants such as Claude Code and "
        "OpenAI Codex CLI to perform deep, physics-based analysis of printed circuit board designs.")

    add_body_text(doc,
        "The framework parses real PCB layout files from major EDA tools (KiCad, Altium, ODB++, Gerber, "
        "IPC-2581) and applies industry-standard formulas and heuristics to identify design issues before "
        "they reach the prototype stage. This \"shift-left\" approach to EMC and SI compliance can save "
        "weeks of redesign time and thousands of dollars in failed compliance testing.")

    add_section_heading(doc, "1.2 Architecture Overview", level=2)
    add_body_text(doc,
        "The system follows a layered architecture with clean separation between protocol handling, data parsing, "
        "analysis computation, and output generation:")
    add_bullet(doc, "MCP Protocol Layer: Exposes 93 tools via the Model Context Protocol standard, enabling "
               "any MCP-compatible AI client to invoke PCB analysis functions.")
    add_bullet(doc, "Session Manager: Maintains parsed design state across multiple tool calls within a review session.")
    add_bullet(doc, "Parser Layer: Supports 8 input formats including KiCad, ODB++, Gerber, IPC-2581, Altium, STEP, "
               "PDF schematics, and BOM files.")
    add_bullet(doc, "Analyzer Layer: 8 analysis domains covering EMC/EMI, signal integrity, power integrity, "
               "thermal, DFM, RF/antenna, ESD, and high-speed digital interfaces.")
    add_bullet(doc, "Report & Visualization Layer: Generates SVG board renders, PNG exports, and professional "
               "DOCX reports with embedded images and findings.")

    add_section_heading(doc, "1.3 Key Value Proposition", level=2)
    add_body_text(doc,
        "MCP PCB EMCopilot provides automated, AI-orchestrated design review that catches EMC, signal integrity, "
        "and manufacturing issues before the first prototype is built. Key benefits include:")
    add_bullet(doc, "Pre-compliance EMC checking: Predict radiated emissions and compare against FCC/CISPR limits "
               "before going to the test lab.")
    add_bullet(doc, "Signal integrity verification: Validate impedance, crosstalk, timing, and eye diagrams for "
               "high-speed interfaces (DDR4/5, PCIe, USB, Ethernet).")
    add_bullet(doc, "Power integrity analysis: Check PDN impedance, decoupling strategy, and VRM placement.")
    add_bullet(doc, "Design for manufacturing: Verify solder paste ratios, component placement, assembly constraints, "
               "and tolerance stackups.")
    add_bullet(doc, "Automated report generation: Professional DOCX reports with annotated board images, "
               "severity-ranked findings, and actionable recommendations.")


def section_architecture_diagram(doc, diagrams):
    add_section_heading(doc, "2. Architecture Diagrams")

    add_section_heading(doc, "2.1 System Architecture", level=2)
    add_body_text(doc,
        "The following diagram shows the end-to-end architecture of the MCP PCB EMCopilot framework, "
        "from MCP client interaction through parsing, analysis, and report generation.")
    if "architecture" in diagrams:
        add_image_from_bytes(doc, diagrams["architecture"],
                            "Figure 1: MCP PCB EMCopilot System Architecture", 6.5)

    add_section_heading(doc, "2.2 PCB Stackup Cross-Section", level=2)
    add_body_text(doc,
        "The framework includes stackup analysis capabilities. Below is a sample 6-layer stackup "
        "showing the layer arrangement, copper weights, and dielectric materials that the tool can parse "
        "and analyze from design files.")
    if "stackup" in diagrams:
        add_image_from_bytes(doc, diagrams["stackup"],
                            "Figure 2: Sample 6-Layer PCB Stackup Cross-Section", 5.5)

    add_section_heading(doc, "2.3 EMI Analysis Flow", level=2)
    add_body_text(doc,
        "EMI analysis follows a signal-chain approach, modeling the path from clock/signal sources "
        "through PCB traces (acting as transmission lines) to unintentional antenna structures, "
        "predicting radiated emissions and comparing against regulatory limits.")
    if "rf_analysis" in diagrams:
        add_image_from_bytes(doc, diagrams["rf_analysis"],
                            "Figure 3: EMI Analysis Signal Flow", 6.0)

    add_section_heading(doc, "2.4 EMC Test Setup Reference", level=2)
    add_body_text(doc,
        "The compliance prediction tools model standard EMC test configurations. Below is a CISPR 25 "
        "radiated emissions test setup, which is the reference geometry used for automotive EMC predictions.")
    if "emc_test" in diagrams:
        add_image_from_bytes(doc, diagrams["emc_test"],
                            "Figure 4: CISPR 25 Radiated Emissions Test Setup", 5.5)


# ---------------------------------------------------------------------------
# Section: Tool Categories
# ---------------------------------------------------------------------------

def section_tool_categories(doc):
    add_section_heading(doc, "3. Tool Categories (93 Tools)")

    add_body_text(doc,
        "The 93 MCP tools are organized into functional categories. Each tool accepts JSON parameters "
        "via the MCP protocol and returns structured JSON results.")

    # Parsers & Session Management
    add_section_heading(doc, "3.1 Parsers & Session Management (14 tools)", level=2)
    add_body_text(doc,
        "These tools handle loading PCB design files and extracting structured data. The parse_layout "
        "tool returns a session_id that subsequent tools use to query the parsed design.")
    tools_parsers = [
        ("pcb_parse_layout", "Parse PCB layout file (KiCad, ODB++, Gerber, IPC-2581, Altium)", "file_path"),
        ("pcb_parse_schematic_pdf", "Parse PDF schematic for components and net labels", "file_path, session_id"),
        ("pcb_parse_step", "Parse STEP file for 3D mechanical review", "file_path"),
        ("pcb_get_components", "Get component list from parsed design", "session_id, filter"),
        ("pcb_get_nets", "Get net list from parsed design", "session_id, filter"),
        ("pcb_get_traces", "Get trace summary from parsed design", "session_id, net_name, layer"),
        ("pcb_get_vias", "Get via list from parsed design", "session_id"),
        ("pcb_get_board_outline", "Get board outline dimensions and area", "session_id"),
        ("pcb_get_stackup", "Get layer stackup from parsed design", "session_id"),
        ("pcb_get_copper_pours", "Get copper pour/zone data per layer", "session_id"),
        ("pcb_get_drill_table", "Get drill table with sizes and counts", "session_id"),
        ("pcb_get_design_rules", "Get extracted DRC constraints", "session_id"),
        ("pcb_list_sessions", "List all active design sessions", ""),
        ("pcb_close_session", "Close a design session and free memory", "session_id"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description", "Key Parameters"],
        tools_parsers,
        col_widths=[2.0, 3.0, 1.8])

    # Impedance Calculators
    add_section_heading(doc, "3.2 Impedance Calculators (7 tools)", level=2)
    add_body_text(doc,
        "Physics-based impedance and transmission line calculators using IPC-2141 formulas "
        "(Hammerstad & Jensen for microstrip, Cohn for stripline).")
    tools_impedance = [
        ("pcb_calc_microstrip_impedance", "Surface trace impedance (IPC-2141)", "width, height, thickness, Er"),
        ("pcb_calc_stripline_impedance", "Buried trace impedance (Cohn)", "width, height, thickness, Er"),
        ("pcb_calc_differential_impedance", "Differential pair impedance", "width, spacing, height, type"),
        ("pcb_calc_cpw_impedance", "Coplanar waveguide impedance", "width, gap, height, Er"),
        ("pcb_calc_trace_width", "Current capacity (IPC-2221)", "current, temp_rise, copper_oz"),
        ("pcb_calc_via_stitching", "Via stitching density for EMI", "frequency, Er"),
        ("pcb_calc_pdn_impedance", "PDN impedance frequency sweep", "VRM, caps, planes, target_Z"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description", "Key Parameters"],
        tools_impedance,
        col_widths=[2.2, 2.5, 2.1])

    # Signal Integrity
    add_section_heading(doc, "3.3 Signal Integrity (14 tools)", level=2)
    add_body_text(doc,
        "Tools for analyzing signal quality, timing margins, crosstalk coupling, via transitions, "
        "and channel loss characteristics.")
    tools_si = [
        ("pcb_analyze_timing", "Analyze setup/hold margins"),
        ("pcb_analyze_crosstalk", "Estimate NEXT and FEXT between traces"),
        ("pcb_analyze_via", "Via inductance, capacitance, impedance"),
        ("pcb_analyze_differential_pair", "Differential pair routing quality"),
        ("pcb_analyze_length_matching", "Trace length matching for signal groups"),
        ("pcb_analyze_mode_conversion", "Differential mode conversion / SCD21"),
        ("pcb_analyze_return_paths", "Return path analysis for all HF nets"),
        ("pcb_analyze_return_current", "Return current density profile"),
        ("pcb_analyze_return_current_density", "Return current density on ref plane"),
        ("pcb_calc_insertion_loss", "Frequency-swept S21 and S11"),
        ("pcb_calc_return_loss", "Return loss, VSWR from mismatch"),
        ("pcb_calc_skin_effect", "Skin depth and AC resistance"),
        ("pcb_calc_dielectric_loss", "Dielectric loss (dB/inch)"),
        ("pcb_calc_eye_diagram", "Statistical eye diagram estimation"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_si,
        col_widths=[2.5, 4.3])

    # EMC / EMI
    add_section_heading(doc, "3.4 EMC / EMI Analysis (14 tools)", level=2)
    add_body_text(doc,
        "Electromagnetic compatibility and interference analysis tools for predicting emissions, "
        "analyzing shielding effectiveness, and checking compliance with regulatory standards.")
    tools_emc = [
        ("pcb_analyze_current_loop", "Radiated emissions from current loops"),
        ("pcb_analyze_clock_emi", "Clock harmonic EMI envelope (trapezoidal)"),
        ("pcb_analyze_smps_emi", "SMPS switching harmonic EMI"),
        ("pcb_analyze_emi_risk", "Per-net EMI risk scoring"),
        ("pcb_analyze_shielding", "Enclosure shielding effectiveness"),
        ("pcb_analyze_grounding", "Grounding topology analysis"),
        ("pcb_analyze_ground_stitch", "Ground via stitching optimization"),
        ("pcb_analyze_common_mode", "Common-mode noise on diff pairs"),
        ("pcb_analyze_cable_coupling", "Cable-to-cable coupling for EMI"),
        ("pcb_analyze_slot_antenna", "Ground plane slot as antenna"),
        ("pcb_analyze_trace_antenna", "Trace as unintentional antenna"),
        ("pcb_estimate_bandwidth", "Signal BW and EMC from rise time"),
        ("pcb_predict_emissions", "Radiated emission spectrum vs limits"),
        ("pcb_predict_compliance", "EMC compliance prediction"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_emc,
        col_widths=[2.5, 4.3])

    # Power Integrity
    add_section_heading(doc, "3.5 Power Integrity (5 tools)", level=2)
    add_body_text(doc,
        "Power distribution network analysis including impedance profiling, decoupling "
        "capacitor placement, and VRM evaluation.")
    tools_pi = [
        ("pcb_analyze_pdn", "PDN impedance analysis"),
        ("pcb_analyze_decoupling", "Decoupling capacitor placement"),
        ("pcb_analyze_vrm", "VRM placement and routing analysis"),
        ("pcb_analyze_copper_spreading", "Copper area heat spreading"),
        ("pcb_calc_plane_resonance", "Power/ground plane resonance"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_pi,
        col_widths=[2.5, 4.3])

    # High-Speed Digital
    add_section_heading(doc, "3.6 High-Speed Digital (8 tools)", level=2)
    add_body_text(doc,
        "Interface-specific analyzers for DDR memory, PCIe, USB, and Ethernet with "
        "protocol-aware timing and topology validation.")
    tools_hs = [
        ("pcb_analyze_ddr", "DDR memory interface routing analysis"),
        ("pcb_analyze_ddr_timing_budget", "Per-lane DDR timing margin (JEDEC)"),
        ("pcb_validate_ddr_topology", "DDR topology auto-detect and validate"),
        ("pcb_analyze_usb", "USB routing analysis"),
        ("pcb_analyze_ethernet", "Ethernet PHY routing analysis"),
        ("pcb_analyze_pcie", "PCIe lane routing analysis"),
        ("pcb_validate_pcie_lanes", "PCIe lane skew validation"),
        ("pcb_calc_pcie_link_budget", "PCIe insertion loss budget"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_hs,
        col_widths=[2.5, 4.3])

    # Thermal
    add_section_heading(doc, "3.7 Thermal Analysis (2 tools)", level=2)
    tools_thermal = [
        ("pcb_analyze_thermal", "Component thermal dissipation analysis"),
        ("pcb_analyze_thermal_via", "Thermal via array analysis"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_thermal,
        col_widths=[2.5, 4.3])

    # DFM / Manufacturing
    add_section_heading(doc, "3.8 DFM / Manufacturing (5 tools)", level=2)
    tools_dfm = [
        ("pcb_analyze_placement", "Component placement manufacturability"),
        ("pcb_analyze_assembly", "Board assembly process considerations"),
        ("pcb_analyze_solder_paste", "Solder paste stencil design analysis"),
        ("pcb_analyze_tolerance", "Manufacturing tolerance stackup"),
        ("pcb_get_manufacturing_notes", "Fab notes and material specs"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_dfm,
        col_widths=[2.5, 4.3])

    # ESD
    add_section_heading(doc, "3.9 ESD Protection (1 tool)", level=2)
    tools_esd = [
        ("pcb_analyze_esd", "ESD protection circuit analysis"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_esd,
        col_widths=[2.5, 4.3])

    # Classification & Detection
    add_section_heading(doc, "3.10 Classification & Detection (4 tools)", level=2)
    add_body_text(doc,
        "Automated classification of nets, interfaces, and overall design type to drive "
        "context-aware analysis selection.")
    tools_class = [
        ("pcb_classify_design", "Classify design type with complexity score"),
        ("pcb_classify_nets", "Classify all nets by function with confidence"),
        ("pcb_detect_interfaces", "Detect high-speed interfaces and pin counts"),
        ("pcb_cross_reference_schematic", "Cross-reference schematic vs layout"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_class,
        col_widths=[2.5, 4.3])

    # Visualization & Rendering
    add_section_heading(doc, "3.11 Visualization & Rendering (5 tools)", level=2)
    tools_vis = [
        ("pcb_render_board", "SVG board view with components, traces, vias"),
        ("pcb_render_net", "SVG highlighting a specific net"),
        ("pcb_render_stackup", "SVG cross-section of layer stackup"),
        ("pcb_annotate_board", "SVG with annotation overlays"),
        ("pcb_get_emi_hotspots", "Identify high-EMI-risk board regions"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_vis,
        col_widths=[2.5, 4.3])

    # Export & Reporting
    add_section_heading(doc, "3.12 Export & Reporting (6 tools)", level=2)
    tools_export = [
        ("pcb_export_render_png", "Export SVG render to PNG file"),
        ("pcb_export_all_renders", "Export all standard renders as PNGs"),
        ("pcb_generate_report", "Structured report from review results"),
        ("pcb_generate_docx_report", "Professional DOCX with embedded images"),
        ("pcb_get_schematic_page", "Get schematic page text and annotations"),
        ("pcb_set_review_context", "Set design review context and requirements"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_export,
        col_widths=[2.5, 4.3])

    # 3D & Enclosure
    add_section_heading(doc, "3.13 3D & Enclosure (3 tools)", level=2)
    tools_3d = [
        ("pcb_get_3d_clearances", "Component-to-component 3D clearances"),
        ("pcb_check_enclosure_fit", "PCB assembly fits within enclosure"),
        ("pcb_find_split_crossings", "Signals crossing ground plane splits"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_3d,
        col_widths=[2.5, 4.3])

    # Orchestration
    add_section_heading(doc, "3.14 Orchestration (2 tools)", level=2)
    add_body_text(doc,
        "High-level tools that coordinate multiple analyzers for comprehensive design review.")
    tools_orch = [
        ("pcb_run_design_review", "Full automated multi-domain design review"),
        ("pcb_optimize_ground_stitching", "Optimize ground via stitching pattern"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_orch,
        col_widths=[2.5, 4.3])

    # Reference Data
    add_section_heading(doc, "3.15 Reference Data & Utilities (3 tools)", level=2)
    tools_ref = [
        ("pcb_get_stackup_templates", "Common PCB stackup configurations"),
        ("pcb_get_material_properties", "Dielectric properties for PCB materials"),
        ("pcb_trace_return_path", "Trace return current path for a specific net"),
    ]
    add_styled_table(doc,
        ["Tool Name", "Description"],
        tools_ref,
        col_widths=[2.5, 4.3])


# ---------------------------------------------------------------------------
# Section: Design Review Workflow
# ---------------------------------------------------------------------------

def section_workflow(doc, diagrams):
    add_section_heading(doc, "4. Design Review Workflow")

    add_body_text(doc,
        "The automated design review process follows a structured pipeline that transforms raw "
        "PCB layout files into actionable engineering findings. The AI orchestrator (Claude Code "
        "or Codex CLI) drives this workflow by calling the appropriate MCP tools in sequence.")

    if "workflow" in diagrams:
        add_image_from_bytes(doc, diagrams["workflow"],
                            "Figure 5: Automated Design Review Workflow", 6.5)

    add_section_heading(doc, "4.1 Step 1: Parse Layout", level=2)
    add_body_text(doc,
        "The review begins with pcb_parse_layout, which accepts KiCad (.kicad_pcb), ODB++ (.zip/.tgz), "
        "Gerber (RS-274X), IPC-2581 (.xml), or Altium files. The parser extracts components, nets, "
        "traces, vias, board outline, stackup, and copper pours into a unified data model. A session_id "
        "is returned for subsequent queries.")

    add_section_heading(doc, "4.2 Step 2: Classify Design", level=2)
    add_body_text(doc,
        "pcb_classify_design analyzes the parsed data to determine the design type "
        "(RF, mixed-signal, high-speed digital, power, or simple digital) and assigns a complexity "
        "score from 1-10. This classification drives which analyzers are selected in Step 4.")

    add_section_heading(doc, "4.3 Step 3: Detect Interfaces", level=2)
    add_body_text(doc,
        "pcb_detect_interfaces identifies high-speed interfaces present in the design (DDR, PCIe, "
        "USB, Ethernet, LVDS, RF) by analyzing net names, component packages, and pin counts. "
        "pcb_classify_nets categorizes every net by function (power, ground, clock, data, analog, etc.).")

    add_section_heading(doc, "4.4 Step 4: Run Domain-Specific Analyzers", level=2)
    add_body_text(doc,
        "Based on the design classification and detected interfaces, the orchestrator selects the "
        "relevant analyzers. For a mixed-signal board with DDR4 and USB, this might include:")
    add_bullet(doc, "EMC: clock EMI, current loop analysis, EMI risk scoring")
    add_bullet(doc, "SI: DDR timing budget, USB differential pair, crosstalk, via analysis")
    add_bullet(doc, "PI: PDN impedance, decoupling, VRM analysis")
    add_bullet(doc, "DFM: placement, solder paste, assembly checks")
    add_bullet(doc, "Thermal: power dissipation, thermal via analysis")

    add_section_heading(doc, "4.5 Step 5: Generate Visualizations", level=2)
    add_body_text(doc,
        "The visualization layer produces SVG renders of the board layout, individual nets, "
        "stackup cross-sections, and annotated boards highlighting findings. These are exported "
        "as PNG images for embedding in reports.")

    add_section_heading(doc, "4.6 Step 6: Export Report", level=2)
    add_body_text(doc,
        "pcb_generate_docx_report creates a professional Word document with embedded board "
        "renders, severity-ranked findings organized by domain, and actionable recommendations. "
        "The report follows a standard template with executive summary, detailed findings, "
        "and appendices.")


# ---------------------------------------------------------------------------
# Section: Supported File Formats
# ---------------------------------------------------------------------------

def section_file_formats(doc):
    add_section_heading(doc, "5. Supported File Formats")

    add_section_heading(doc, "5.1 Input Formats", level=2)
    formats_in = [
        ("KiCad", ".kicad_pcb", "Full layout data: components, nets, traces, vias, zones, stackup, board outline"),
        ("ODB++", ".zip / .tgz", "Complete manufacturing data: all layers, drill data, netlist, components"),
        ("Gerber", "RS-274X", "Layer-by-layer artwork: copper, mask, silk, drill (Excellon)"),
        ("IPC-2581", ".xml", "Industry standard data exchange format with full design intent"),
        ("Altium", ".PcbDoc", "Altium Designer binary format (requires olefile)"),
        ("STEP", ".step / .stp", "3D mechanical model for clearance and enclosure analysis"),
        ("PDF Schematic", ".pdf", "Schematic pages with text-layer extraction (requires pymupdf)"),
        ("BOM", ".csv / .xlsx", "Bill of materials for component validation"),
    ]
    add_styled_table(doc,
        ["Format", "Extensions", "Data Extracted"],
        formats_in,
        col_widths=[1.2, 1.2, 4.4])

    add_section_heading(doc, "5.2 Output Formats", level=2)
    formats_out = [
        ("SVG", "Board renders, net highlights, stackup cross-sections, annotated boards"),
        ("PNG", "Exported from SVG renders via cairosvg for embedding in documents"),
        ("DOCX", "Professional design review reports with embedded images and structured findings"),
        ("JSON", "Structured analysis results returned via MCP protocol"),
    ]
    add_styled_table(doc,
        ["Format", "Usage"],
        formats_out,
        col_widths=[1.2, 5.6])


# ---------------------------------------------------------------------------
# Section: Simulation & Analysis Capabilities
# ---------------------------------------------------------------------------

def section_analysis_capabilities(doc):
    add_section_heading(doc, "6. Simulation & Analysis Capabilities")

    add_body_text(doc,
        "The framework implements physics-based calculations grounded in industry standards "
        "and peer-reviewed electromagnetic theory. The following subsections detail the key "
        "computational methods.")

    add_section_heading(doc, "6.1 Impedance Calculations (IPC-2141)", level=2)
    add_body_text(doc,
        "Microstrip impedance uses the Hammerstad & Jensen formulation with thickness correction. "
        "Stripline impedance uses the Cohn formula. Both account for effective dielectric constant "
        "and conductor thickness effects.")
    formulas = [
        ("Microstrip", "Hammerstad & Jensen (IPC-2141)", "Effective width correction for trace thickness"),
        ("Stripline", "Cohn (IPC-2141)", "Symmetric or offset stripline between planes"),
        ("Differential Pair", "Coupled line theory", "Coupling factor from trace spacing / height ratio"),
        ("CPW / GCPW", "Conformal mapping", "Grounded and ungrounded coplanar waveguide"),
    ]
    add_styled_table(doc,
        ["Trace Type", "Formula Basis", "Notes"],
        formulas,
        col_widths=[1.5, 2.2, 3.1])

    add_section_heading(doc, "6.2 Current Capacity (IPC-2221)", level=2)
    add_body_text(doc,
        "Trace width calculations use the IPC-2221 current-temperature rise relationship "
        "with separate coefficients for internal and external layers. The formula accounts for "
        "copper weight, ambient temperature, and maximum allowable temperature rise.")

    add_section_heading(doc, "6.3 EMI Prediction", level=2)
    add_body_text(doc,
        "Clock and signal EMI prediction uses trapezoidal waveform Fourier analysis to calculate "
        "harmonic content. The envelope of harmonics follows the well-known 0 dB/decade, "
        "-20 dB/decade, and -40 dB/decade slopes defined by the fundamental frequency, "
        "1/(pi * tau_r) and 1/(pi * tau_f) corner frequencies. Emissions are calculated using "
        "the magnetic dipole antenna model for PCB current loops.")

    add_section_heading(doc, "6.4 Eye Diagram Estimation", level=2)
    add_body_text(doc,
        "Statistical eye diagram calculation models the channel as a lossy transmission line "
        "with frequency-dependent conductor loss (skin effect + Hammerstad roughness correction) "
        "and dielectric loss. Eye height and width are estimated from the channel impulse response "
        "and compared against standard thresholds for the target data rate.")

    add_section_heading(doc, "6.5 PDN Impedance", level=2)
    add_body_text(doc,
        "PDN impedance profiling models the VRM output impedance, bulk capacitors, MLCC decoupling "
        "capacitors, and plane capacitance as parallel RLC networks. The frequency sweep identifies "
        "anti-resonance peaks that exceed the target impedance Z_target = V * ripple% / I_max.")

    add_section_heading(doc, "6.6 Thermal Analysis", level=2)
    add_body_text(doc,
        "Thermal resistance chain calculation models junction-to-ambient paths through the "
        "die, package, PCB copper, thermal vias, and convection. Thermal via arrays are analyzed "
        "for effective thermal conductivity enhancement.")

    add_section_heading(doc, "6.7 Via Modeling", level=2)
    add_body_text(doc,
        "Via electrical models calculate parasitic inductance (from via barrel length and diameter), "
        "capacitance (from pad-to-antipad geometry), and characteristic impedance. These are used "
        "for signal integrity assessment of layer transitions.")

    add_section_heading(doc, "6.8 Crosstalk Analysis", level=2)
    add_body_text(doc,
        "Near-end crosstalk (NEXT) and far-end crosstalk (FEXT) coupling coefficients are calculated "
        "from trace geometry (width, spacing, dielectric height) using coupled microstrip/stripline "
        "theory. The analysis accounts for coupling length, rise time, and termination conditions.")


# ---------------------------------------------------------------------------
# Section: Integration with Other MCP Servers
# ---------------------------------------------------------------------------

def section_integrations(doc):
    add_section_heading(doc, "7. Integration with Other MCP Servers")

    add_body_text(doc,
        "MCP PCB EMCopilot is designed to work alongside other specialized MCP servers for "
        "a complete PCB engineering workflow. The AI orchestrator can invoke tools across "
        "multiple MCP servers in a single design review session.")

    integrations = [
        ("mcp-emc-regulations",
         "FCC Part 15, CISPR 11/22/25/32, IEC 61000-4-x",
         "Provides regulatory emission limits and immunity test levels that EMCopilot's "
         "compliance prediction tools compare against. Enables automated pass/fail checking "
         "with exact limit values at each frequency."),
        ("mcp-nec2-antenna",
         "NEC2 engine for dipole, yagi, loop, vertical",
         "Simulates unintentional antenna structures identified by EMCopilot's trace and slot "
         "antenna analyzers. Provides radiation pattern and gain data for more accurate "
         "emission predictions."),
        ("mcp-openems",
         "Full-wave EM simulation via openEMS",
         "Validates EMCopilot's approximate calculations with full 3D electromagnetic simulation. "
         "Useful for complex structures like via transitions, plane resonances, and connector launches."),
        ("mcp-drawio-engineering",
         "Engineering diagram generation",
         "Generates PCB stackup diagrams, RF block diagrams, and EMC test setup illustrations. "
         "Used to create visual documentation for design review reports."),
    ]

    for name, capability, description in integrations:
        add_section_heading(doc, name, level=2)
        add_body_text(doc, f"Capability: {capability}")
        add_body_text(doc, description)


# ---------------------------------------------------------------------------
# Section: Getting Started
# ---------------------------------------------------------------------------

def section_getting_started(doc):
    add_section_heading(doc, "8. Getting Started")

    add_section_heading(doc, "8.1 Installation", level=2)
    add_body_text(doc, "Clone the repository and install with uv:")
    add_code_block(doc,
        "git clone https://github.com/RFingAdam/mcp-pcb-emcopilot.git\n"
        "cd mcp-pcb-emcopilot\n"
        "uv pip install -e .")
    add_body_text(doc, "For full feature support, install optional dependencies:")
    add_code_block(doc, "uv pip install -e '.[all]'")
    add_body_text(doc, "This installs: olefile (Altium), openpyxl (Excel BOM), numpy (S-parameters), "
                  "networkx (ground analysis), pymupdf (PDF schematics).")

    add_section_heading(doc, "8.2 Configuration", level=2)
    add_body_text(doc, "Add to Claude Code:")
    add_code_block(doc,
        "claude mcp add pcb-emcopilot -- uv run --directory /path/to/mcp-pcb-emcopilot mcp-pcb-emcopilot")
    add_body_text(doc, "Add to Codex CLI:")
    add_code_block(doc,
        "codex mcp add pcb-emcopilot -- uv run --directory /path/to/mcp-pcb-emcopilot mcp-pcb-emcopilot")
    add_body_text(doc, "Or add to MCP client config file (JSON):")
    add_code_block(doc,
        '{\n'
        '  "command": "uv",\n'
        '  "args": ["run", "--directory", "/path/to/mcp-pcb-emcopilot", "mcp-pcb-emcopilot"]\n'
        '}')

    add_section_heading(doc, "8.3 Example Workflow", level=2)
    add_body_text(doc, "A typical design review session with Claude Code:")

    add_body_text(doc, "Step 1: Parse the PCB layout")
    add_code_block(doc,
        'User: "Parse my KiCad board at /designs/my_board.kicad_pcb and run a full design review."')
    add_body_text(doc, "Step 2: Claude Code calls the tools")
    add_code_block(doc,
        "# AI orchestrator calls these tools automatically:\n"
        "pcb_parse_layout(file_path='/designs/my_board.kicad_pcb')\n"
        "pcb_classify_design(session_id='abc123')\n"
        "pcb_detect_interfaces(session_id='abc123')\n"
        "pcb_run_design_review(session_id='abc123')\n"
        "pcb_render_board(session_id='abc123')\n"
        "pcb_export_all_renders(session_id='abc123')\n"
        "pcb_generate_docx_report(session_id='abc123', output_path='review.docx')")
    add_body_text(doc, "Step 3: Review the output")
    add_code_block(doc,
        "# The AI returns a structured summary of findings:\n"
        "# - 2 CRITICAL: DDR4 trace impedance mismatch, missing return via\n"
        "# - 5 WARNING: Crosstalk on CLK net, via stub resonance, ...\n"
        "# - 8 INFO: Copper pour coverage, stackup recommendations, ...\n"
        "# Full report saved to review.docx")


# ---------------------------------------------------------------------------
# Section: API Reference Summary
# ---------------------------------------------------------------------------

def section_api_reference(doc):
    add_section_heading(doc, "9. API Reference Summary")

    add_body_text(doc,
        "This section provides a condensed reference for all 93 tools organized alphabetically "
        "within each category. For complete parameter documentation, refer to the tool descriptions "
        "exposed via the MCP protocol (visible in your MCP client's tool listing).")

    # Comprehensive quick reference
    all_tools = [
        # Parsers
        ("pcb_parse_layout", "Parse PCB layout file", "file_path: str"),
        ("pcb_parse_schematic_pdf", "Parse PDF schematic", "file_path: str, session_id?: str"),
        ("pcb_parse_step", "Parse STEP 3D file", "file_path: str"),
        ("pcb_get_board_outline", "Board dimensions", "session_id: str"),
        ("pcb_get_components", "Component list", "session_id: str, filter?: str"),
        ("pcb_get_copper_pours", "Copper pour data", "session_id: str"),
        ("pcb_get_design_rules", "DRC constraints", "session_id: str"),
        ("pcb_get_drill_table", "Drill sizes/counts", "session_id: str"),
        ("pcb_get_manufacturing_notes", "Fab notes", "session_id: str"),
        ("pcb_get_nets", "Net list", "session_id: str, filter?: str"),
        ("pcb_get_stackup", "Layer stackup", "session_id: str"),
        ("pcb_get_traces", "Trace summary", "session_id: str, net_name?: str"),
        ("pcb_get_vias", "Via list", "session_id: str"),
        ("pcb_list_sessions", "Active sessions", ""),
        ("pcb_close_session", "Close session", "session_id: str"),
        # Impedance
        ("pcb_calc_cpw_impedance", "CPW impedance", "width, gap, height, Er, thickness"),
        ("pcb_calc_differential_impedance", "Diff pair impedance", "width, spacing, height, thickness, Er, type"),
        ("pcb_calc_microstrip_impedance", "Microstrip impedance", "width, height, thickness, Er"),
        ("pcb_calc_stripline_impedance", "Stripline impedance", "width, height, thickness, Er"),
        ("pcb_calc_trace_width", "Trace current capacity", "current, temp_rise, copper_oz, layer"),
        ("pcb_calc_via_stitching", "Via stitch spacing", "frequency, Er"),
        ("pcb_calc_pdn_impedance", "PDN impedance sweep", "VRM, caps, planes, freq range"),
        # SI
        ("pcb_analyze_crosstalk", "NEXT/FEXT coupling", "width, spacing, height, length, rise_time"),
        ("pcb_analyze_differential_pair", "Diff pair quality", "session_id, net_p, net_n"),
        ("pcb_analyze_length_matching", "Length matching", "session_id, net_names, tolerance"),
        ("pcb_analyze_mode_conversion", "Mode conversion", "Z_diff, Z_comm, length, freq"),
        ("pcb_analyze_return_current", "Return current profile", "width, height, frequency"),
        ("pcb_analyze_return_current_density", "Return density map", "session_id, net_name"),
        ("pcb_analyze_return_paths", "Return path analysis", "session_id"),
        ("pcb_analyze_timing", "Timing margins", "freq, length, setup, hold, rise_time"),
        ("pcb_analyze_via", "Via characteristics", "drill, pad, antipad, length"),
        ("pcb_calc_dielectric_loss", "Dielectric loss", "freq, Er, Df, length"),
        ("pcb_calc_eye_diagram", "Eye diagram", "data_rate, length, Er, Df, width, height"),
        ("pcb_calc_insertion_loss", "Insertion loss S21", "freq, length, width, height, Er, Df"),
        ("pcb_calc_return_loss", "Return loss S11", "Z_line, Z_load"),
        ("pcb_calc_skin_effect", "Skin depth/loss", "freq, conductivity, roughness"),
        # EMC
        ("pcb_analyze_cable_coupling", "Cable coupling", "length, spacing, freq, type"),
        ("pcb_analyze_clock_emi", "Clock EMI harmonics", "freq, amplitude, rise_time, loop_area"),
        ("pcb_analyze_common_mode", "Common-mode noise", "Z_diff, skew, freq"),
        ("pcb_analyze_current_loop", "Loop radiation", "freq, current, area"),
        ("pcb_analyze_emi_risk", "EMI risk scoring", "session_id"),
        ("pcb_analyze_grounding", "Grounding topology", "session_id"),
        ("pcb_analyze_ground_stitch", "Via stitch spacing", "frequency, lambda_fraction"),
        ("pcb_analyze_shielding", "Shielding SE", "material, thickness, freq, aperture"),
        ("pcb_analyze_slot_antenna", "Slot as antenna", "length, width, freq"),
        ("pcb_analyze_smps_emi", "SMPS EMI", "freq, duty, current, loop_area"),
        ("pcb_analyze_trace_antenna", "Trace as antenna", "length, freq"),
        ("pcb_estimate_bandwidth", "BW from rise time", "rise_time"),
        ("pcb_predict_compliance", "Compliance prediction", "session_id, standard"),
        ("pcb_predict_emissions", "Emission spectrum", "session_id, limit_standard"),
        # PI
        ("pcb_analyze_copper_spreading", "Cu heat spreading", "area, thickness, power"),
        ("pcb_analyze_decoupling", "Decap placement", "session_id, ic_name"),
        ("pcb_analyze_pdn", "PDN impedance", "session_id, rail_name"),
        ("pcb_analyze_vrm", "VRM analysis", "session_id"),
        ("pcb_calc_plane_resonance", "Plane resonance", "length, width, Er, height"),
        # High-Speed
        ("pcb_analyze_ddr", "DDR interface", "session_id, ddr_type"),
        ("pcb_analyze_ddr_timing_budget", "DDR timing budget", "session_id, ddr_type, data_rate"),
        ("pcb_analyze_ethernet", "Ethernet routing", "session_id"),
        ("pcb_analyze_pcie", "PCIe lanes", "session_id, gen"),
        ("pcb_analyze_usb", "USB routing", "session_id, usb_version"),
        ("pcb_calc_pcie_link_budget", "PCIe link budget", "gen, trace_len, via_count"),
        ("pcb_validate_ddr_topology", "DDR topology", "session_id"),
        ("pcb_validate_pcie_lanes", "PCIe lane skew", "session_id, gen"),
        # Thermal
        ("pcb_analyze_thermal", "Thermal dissipation", "power, theta_ja, T_ambient"),
        ("pcb_analyze_thermal_via", "Thermal via array", "via_count, drill, length"),
        # DFM
        ("pcb_analyze_assembly", "Assembly check", "session_id"),
        ("pcb_analyze_placement", "Placement check", "session_id"),
        ("pcb_analyze_solder_paste", "Solder paste", "session_id"),
        ("pcb_analyze_tolerance", "Tolerance stackup", "dimensions, tolerances"),
        # ESD
        ("pcb_analyze_esd", "ESD protection", "session_id, port_name"),
        # Classification
        ("pcb_classify_design", "Design classification", "session_id"),
        ("pcb_classify_nets", "Net classification", "session_id"),
        ("pcb_cross_reference_schematic", "Schem vs layout", "session_id"),
        ("pcb_detect_interfaces", "Interface detection", "session_id"),
        # Visualization
        ("pcb_annotate_board", "Annotated board SVG", "session_id, annotations"),
        ("pcb_get_emi_hotspots", "EMI hotspot regions", "session_id"),
        ("pcb_render_board", "Board SVG render", "session_id, layers, highlights"),
        ("pcb_render_net", "Net SVG render", "session_id, net_name"),
        ("pcb_render_stackup", "Stackup SVG render", "session_id"),
        # Export
        ("pcb_export_all_renders", "Export all PNGs", "session_id, output_dir"),
        ("pcb_export_render_png", "Export SVG to PNG", "svg_content, output_path"),
        ("pcb_generate_docx_report", "DOCX report", "session_id, output_path"),
        ("pcb_generate_report", "Structured report", "session_id, format"),
        ("pcb_get_schematic_page", "Schematic page data", "session_id, page_number"),
        ("pcb_set_review_context", "Set review context", "session_id, requirements"),
        # 3D
        ("pcb_check_enclosure_fit", "Enclosure fit check", "session_id, enclosure dims"),
        ("pcb_find_split_crossings", "Split plane crossings", "session_id"),
        ("pcb_get_3d_clearances", "3D clearances", "session_id"),
        # Orchestration
        ("pcb_optimize_ground_stitching", "Optimize stitching", "session_id, frequency"),
        ("pcb_run_design_review", "Full design review", "session_id"),
        # Reference
        ("pcb_get_material_properties", "Material database", ""),
        ("pcb_get_stackup_templates", "Stackup templates", ""),
        ("pcb_trace_return_path", "Trace return path", "session_id, net_name"),
    ]

    # Split into pages of reasonable table size
    page_size = 30
    for i in range(0, len(all_tools), page_size):
        chunk = all_tools[i:i + page_size]
        add_styled_table(doc,
            ["Tool Name", "Description", "Key Parameters"],
            chunk,
            col_widths=[2.2, 1.8, 2.8])


# ---------------------------------------------------------------------------
# Section: Common Design Targets
# ---------------------------------------------------------------------------

def section_design_targets(doc):
    add_section_heading(doc, "10. Common Design Targets")

    add_body_text(doc,
        "The following tables provide quick-reference impedance targets and material properties "
        "commonly used with the framework's calculators.")

    add_section_heading(doc, "10.1 Interface Impedance Targets", level=2)
    targets = [
        ("General purpose", "50", "-"),
        ("USB 2.0", "-", "90"),
        ("USB 3.x", "-", "85"),
        ("HDMI", "-", "100"),
        ("DDR4", "40", "80"),
        ("DDR5", "40", "80"),
        ("PCIe Gen3/4", "-", "85"),
        ("PCIe Gen5/6", "-", "85"),
        ("Ethernet 100BASE-TX", "-", "100"),
        ("Ethernet 1000BASE-T", "-", "100"),
        ("LVDS", "-", "100"),
        ("SATA", "-", "85"),
    ]
    add_styled_table(doc,
        ["Interface", "Single-Ended (Ohm)", "Differential (Ohm)"],
        targets,
        col_widths=[2.0, 2.0, 2.0])

    add_section_heading(doc, "10.2 PCB Material Properties", level=2)
    materials = [
        ("FR-4 Standard", "4.2-4.5", "0.020", "140-170"),
        ("FR-4 High-Tg", "4.2-4.4", "0.018", "170-210"),
        ("Rogers RO4003C", "3.38", "0.0027", "280"),
        ("Rogers RO4350B", "3.48", "0.0037", "280"),
        ("Isola I-Speed", "3.63", "0.0082", "200"),
        ("Megtron 6 (R-5775)", "3.4", "0.002", "230"),
        ("Polyimide (Flex)", "3.3-3.5", "0.003", "260+"),
    ]
    add_styled_table(doc,
        ["Material", "Dk (at 1 GHz)", "Df (Loss Tan.)", "Tg (C)"],
        materials,
        col_widths=[2.0, 1.5, 1.5, 1.0])


# ---------------------------------------------------------------------------
# Main document generation
# ---------------------------------------------------------------------------

def generate_document():
    """Generate the complete framework reference document."""
    print("Generating diagrams...")
    diagrams = generate_all_diagrams()

    print("Creating DOCX document...")
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Narrow margins for more content space
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Build document sections
    print("  Building cover page...")
    create_cover_page(doc)

    print("  Building table of contents...")
    add_toc(doc)

    print("  Building framework overview...")
    section_overview(doc)
    doc.add_page_break()

    print("  Building architecture diagrams...")
    section_architecture_diagram(doc, diagrams)
    doc.add_page_break()

    print("  Building tool categories...")
    section_tool_categories(doc)
    doc.add_page_break()

    print("  Building design review workflow...")
    section_workflow(doc, diagrams)
    doc.add_page_break()

    print("  Building file formats section...")
    section_file_formats(doc)
    doc.add_page_break()

    print("  Building analysis capabilities...")
    section_analysis_capabilities(doc)
    doc.add_page_break()

    print("  Building integrations section...")
    section_integrations(doc)
    doc.add_page_break()

    print("  Building getting started section...")
    section_getting_started(doc)
    doc.add_page_break()

    print("  Building API reference...")
    section_api_reference(doc)
    doc.add_page_break()

    print("  Building design targets...")
    section_design_targets(doc)

    # Save
    print(f"\nSaving document to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print(f"Document saved successfully ({os.path.getsize(OUTPUT_PATH)} bytes)")


if __name__ == "__main__":
    generate_document()
