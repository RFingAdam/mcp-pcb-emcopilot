"""Professional DOCX report generator for PCB design reviews.

Generates a multi-section design review document with embedded images
from board renders, schematic pages, and stackup cross-sections.

Optional dependency: python-docx (pip install python-docx).
"""

from __future__ import annotations

import os
import tempfile
from dataclasses import dataclass, field
from typing import Optional

from ..models.pcb_data import PCBDesignData


@dataclass
class ReportImage:
    """An image to embed in the report."""
    path: str
    caption: str
    width_inches: float = 6.0


@dataclass
class ReportFinding:
    """A design review finding."""
    severity: str  # CRITICAL, HIGH, WARNING, PASS, INFO
    title: str
    detail: str
    recommendation: str = ""
    image: Optional[ReportImage] = None


@dataclass
class ReportSection:
    """A section of the report."""
    title: str
    level: int = 1
    text: str = ""
    findings: list[ReportFinding] = field(default_factory=list)
    tables: list[dict] = field(default_factory=list)
    images: list[ReportImage] = field(default_factory=list)
    subsections: list[ReportSection] = field(default_factory=list)


def _check_docx():
    """Verify python-docx is available."""
    try:
        from docx import Document  # noqa: F401
        return True
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX report generation. "
            "Install with: pip install python-docx"
        )


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex: str):
    """Set table cell background colour."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    """Add a professionally styled table to the document."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, header_color)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                _set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    return table


def add_finding_box(doc, severity, title, detail, recommendation=""):
    """Add a colour-coded finding box."""
    from docx.shared import Inches, Pt, RGBColor

    colors = {
        "CRITICAL": ("C00000", "FBE5D6"),
        "HIGH": ("C00000", "FCE4EC"),
        "WARNING": ("7F6000", "FFF8E1"),
        "PASS": ("1B5E20", "E8F5E9"),
        "INFO": ("1F4E79", "E3F2FD"),
    }
    text_color, _ = colors.get(severity, ("333333", "F5F5F5"))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(f"[{severity}] ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(text_color)
    run.font.size = Pt(10)

    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)

    if detail:
        p2 = doc.add_paragraph(detail)
        p2.paragraph_format.left_indent = Inches(0.3)
        p2.runs[0].font.size = Pt(9)
        p2.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    if recommendation:
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Inches(0.3)
        run = p3.add_run("Recommendation: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        run = p3.add_run(recommendation)
        run.font.size = Pt(9)


def add_image_with_caption(doc, path, caption, width_inches=6.0):
    """Add an image with a centred caption."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    if not os.path.exists(path):
        p = doc.add_paragraph(f"[Image not available: {path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(8)
    cap.runs[0].font.italic = True
    cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    cap.paragraph_format.space_after = Pt(12)


# ---------------------------------------------------------------------------
# Report image generation helpers
# ---------------------------------------------------------------------------

def generate_all_renders(
    design: PCBDesignData,
    session_id: str,
    output_dir: str,
    width_px: int = 1600,
) -> dict[str, str]:
    """Generate all board/stackup/net renders and export as PNG.

    Args:
        design: Parsed PCB design data.
        session_id: Active session ID.
        output_dir: Directory to write PNG files.
        width_px: Width of rasterised images.

    Returns:
        Dict mapping render label to output PNG path.
    """
    from ..visualization.annotator import Annotator
    from ..visualization.board_renderer import BoardRenderer
    from ..visualization.exporter import svg_to_png
    from ..visualization.stackup_renderer import StackupRenderer

    os.makedirs(output_dir, exist_ok=True)
    results: dict[str, str] = {}

    # Full board render
    renderer = BoardRenderer(design, width_px=width_px)
    svg = renderer.render_board()
    results["board_full"] = svg_to_png(svg, os.path.join(output_dir, "board_full.png"), width=width_px)

    # Stackup
    su_renderer = StackupRenderer(design)
    svg = su_renderer.render()
    results["stackup"] = svg_to_png(svg, os.path.join(output_dir, "stackup.png"), width=width_px)

    # Classified net highlight renders
    net_groups = _classify_render_nets(design)
    for group_label, nets in net_groups.items():
        try:
            renderer = BoardRenderer(design, width_px=width_px)
            svg = renderer.render_board(highlight_nets=nets)
            out = os.path.join(output_dir, f"nets_{group_label}.png")
            results[f"nets_{group_label}"] = svg_to_png(svg, out, width=width_px)
        except Exception:
            pass

    # Annotated board with findings
    if design.review_results:
        annotations = _build_annotations_from_findings(design)
        if annotations:
            annotator = Annotator(design)
            svg = annotator.render_annotated_board(annotations=annotations)
            out = os.path.join(output_dir, "board_annotated.png")
            results["board_annotated"] = svg_to_png(svg, out, width=width_px)

    # Individual net renders for key nets
    key_nets = _get_key_nets(design)
    for net_name in key_nets[:8]:
        try:
            renderer = BoardRenderer(design)
            svg = renderer.render_net(net_name)
            safe = net_name.replace("/", "_").replace(" ", "_")
            out = os.path.join(output_dir, f"net_{safe}.png")
            results[f"net_{net_name}"] = svg_to_png(svg, out, width=width_px)
        except Exception:
            pass

    return results


def _classify_render_nets(design: PCBDesignData) -> dict[str, list[str]]:
    """Group nets by type for highlight renders."""
    groups: dict[str, list[str]] = {}
    for net in design.nets:
        name = net.get("name", "") if isinstance(net, dict) else getattr(net, "name", "")
        upper = name.upper()
        if "USB" in upper:
            groups.setdefault("usb", []).append(name)
        elif "DDR" in upper:
            groups.setdefault("ddr", []).append(name)
        elif "RF" in upper or "WIFI" in upper or "BLE" in upper or "HALOW" in upper:
            groups.setdefault("rf", []).append(name)
        elif "ETH" in upper or "TX_" in upper or "RX_" in upper:
            groups.setdefault("ethernet", []).append(name)
        elif "BUCK" in upper or "LDO" in upper or "VCC" in upper or "PWR" in upper:
            groups.setdefault("power", []).append(name)

    # Limit to 10 nets per group to keep renders readable
    return {k: v[:10] for k, v in groups.items()}


def _get_key_nets(design: PCBDesignData) -> list[str]:
    """Select representative nets for individual renders."""
    key = []
    seen_types = set()
    for net in design.nets:
        name = net.get("name", "") if isinstance(net, dict) else getattr(net, "name", "")
        upper = name.upper()
        ntype = None
        if "DDR_DQ0" == upper:
            ntype = "ddr"
        elif "GND" == upper:
            ntype = "gnd"
        elif "USB" in upper and "D_P" in upper and "usb" not in seen_types:
            ntype = "usb"
        elif "RF" in upper and "rf" not in seen_types:
            ntype = "rf"
        elif "TX_P" == upper:
            ntype = "eth"
        elif ("BUCK" in upper or "LDO" in upper) and "power" not in seen_types:
            ntype = "power"
        elif "WIFI" in upper and "wifi" not in seen_types:
            ntype = "wifi"

        if ntype and ntype not in seen_types:
            key.append(name)
            seen_types.add(ntype)

    return key


def _build_annotations_from_findings(design: PCBDesignData) -> list[dict]:
    """Create annotation overlays from design review results."""
    annotations = []
    results = design.review_results or {}

    # EMI hotspots
    for hs in results.get("emi_hotspots", []):
        annotations.append({
            "type": "warning",
            "x": hs.get("center_x_mm", 0),
            "y": hs.get("center_y_mm", 0),
            "severity": "error",
            "text": f"EMI Hotspot ({hs.get('risk_score', 0)}/100)",
        })

    # Component-level findings: map IC positions
    comp_map = {}
    for c in design.components:
        ref = c.get("reference", "") if isinstance(c, dict) else getattr(c, "reference", "")
        if ref:
            x = c.get("x_mm", 0) if isinstance(c, dict) else getattr(c, "x_mm", 0)
            y = c.get("y_mm", 0) if isinstance(c, dict) else getattr(c, "y_mm", 0)
            comp_map[ref] = (x, y)

    # Walk domain results for component-related findings
    for dr in results.get("domain_results", []):
        for finding in dr.get("findings", []):
            severity = finding.get("severity", "INFO")
            if severity in ("CRITICAL", "HIGH"):
                text = finding.get("title", finding.get("description", ""))[:50]
                # Try to locate on board
                for ref, (x, y) in comp_map.items():
                    if ref.upper() in text.upper():
                        annotations.append({
                            "type": "warning",
                            "x": x, "y": y,
                            "severity": "error" if severity == "CRITICAL" else "warning",
                            "text": text,
                        })
                        break

    return annotations


# ---------------------------------------------------------------------------
# Main DOCX generation
# ---------------------------------------------------------------------------

def generate_docx_report(
    design: PCBDesignData,
    session_id: str,
    output_path: Optional[str] = None,
    image_dir: Optional[str] = None,
    images: Optional[dict[str, str]] = None,
    title: str = "PCB Design Review Report",
    subtitle: str = "",
    author: str = "MCP PCB EMCopilot",
) -> str:
    """Generate a professional DOCX design review report.

    Args:
        design: PCBDesignData with review_results populated.
        session_id: Session identifier.
        output_path: Destination .docx path. Defaults to temp file.
        image_dir: Directory containing pre-rendered images.
        images: Explicit mapping of label -> image path (overrides image_dir).
        title: Report title.
        subtitle: Report subtitle (e.g. project name).
        author: Author name for cover page.

    Returns:
        Absolute path to the generated DOCX file.
    """
    _check_docx()
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    if output_path is None:
        fd, output_path = tempfile.mkstemp(suffix=".docx", prefix="pcb_review_")
        os.close(fd)

    # Resolve images
    img = images or {}
    if image_dir and os.path.isdir(image_dir):
        for fname in os.listdir(image_dir):
            if fname.endswith((".png", ".jpg", ".jpeg")):
                label = os.path.splitext(fname)[0]
                img.setdefault(label, os.path.join(image_dir, fname))

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ---- Cover page ----
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Comprehensive EMC, Signal Integrity & DFM Analysis")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(4):
        doc.add_paragraph()

    # Cover info
    board_w = design.board_width_mm or 0
    board_h = design.board_height_mm or 0
    layers = len(design.layers) if design.layers else 0
    cover_data = [
        ("Board Size", f"{board_w:.1f} x {board_h:.1f} mm"),
        ("Layer Count", str(layers)),
        ("Components", str(len(design.components))),
        ("Nets", str(len(design.nets))),
        ("Review Tool", f"MCP PCB EMCopilot ({author})"),
    ]
    table = doc.add_table(rows=len(cover_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(cover_data):
        row = table.rows[i]
        row.cells[0].text = ""
        row.cells[1].text = ""
        p0 = row.cells[0].paragraphs[0]
        run = p0.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1 = row.cells[1].paragraphs[0]
        run = p1.add_run(value)
        run.font.size = Pt(10)
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.5)

    doc.add_page_break()

    # ---- Board overview ----
    doc.add_heading("Board Overview & Layout", level=1)
    add_styled_table(doc,
        ["Parameter", "Value"],
        [
            ("Board Size", f"{board_w:.1f} x {board_h:.1f} mm ({board_w * board_h:.0f} mm²)"),
            ("Layer Count", str(layers)),
            ("Components", str(len(design.components))),
            ("Nets", str(len(design.nets))),
            ("Traces", str(len(design.traces))),
            ("Vias", str(len(design.vias))),
            ("Copper Zones", str(len(design.zones))),
        ],
        col_widths=[2.5, 4.0],
    )

    # Board render
    for key in ("board_annotated", "board_full"):
        if key in img:
            add_image_with_caption(doc, img[key],
                f"Board layout — {board_w:.0f} x {board_h:.0f} mm, {layers} layers, "
                f"{len(design.components)} components, {len(design.traces)} traces.",
                width_inches=6.2)
            break

    doc.add_page_break()

    # ---- Stackup ----
    doc.add_heading("Layer Stackup", level=1)
    if "stackup" in img:
        add_image_with_caption(doc, img["stackup"],
            "Layer stackup cross-section showing all copper and dielectric layers.",
            width_inches=5.5)

    if design.layers:
        layer_rows = []
        for ly in design.layers:
            lname = ly.get("name", "") if isinstance(ly, dict) else getattr(ly, "name", "")
            ltype = ly.get("type", "") if isinstance(ly, dict) else getattr(ly, "type", "")
            layer_rows.append((lname, ltype))
        if layer_rows:
            add_styled_table(doc, ["Layer", "Type"], layer_rows, col_widths=[3.0, 3.5])

    doc.add_page_break()

    # ---- Design review results ----
    results = design.review_results or {}
    summary = results.get("executive_summary", {})

    if summary:
        doc.add_heading("Executive Summary", level=1)
        status = summary.get("overall_status", "UNKNOWN")
        p = doc.add_paragraph()
        run = p.add_run(f"Overall Assessment: {status}")
        run.bold = True
        run.font.size = Pt(12)
        if "CRITICAL" in status.upper() or "FAIL" in status.upper():
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif "PASS" in status.upper():
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        else:
            run.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)

        summary_rows = [
            ("Design Complexity", str(summary.get("complexity", "—"))),
            ("Total Findings", str(summary.get("total_findings", 0))),
            ("Critical", str(summary.get("total_critical", 0))),
            ("Warnings", str(summary.get("total_warnings", 0))),
            ("Domains Analyzed", str(summary.get("domains_analyzed", 0))),
        ]
        add_styled_table(doc, ["Metric", "Value"], summary_rows, col_widths=[3.0, 3.5])
        doc.add_page_break()

    # ---- Domain results ----
    figure_num = 1
    for dr in results.get("domain_results", []):
        domain = dr.get("domain", "Unknown")
        doc.add_heading(f"{domain} Analysis", level=1)

        status = dr.get("status", "")
        if status:
            p = doc.add_paragraph()
            run = p.add_run(f"Status: {status}")
            run.bold = True
            run.font.size = Pt(10)

        for finding in dr.get("findings", []):
            sev = finding.get("severity", "INFO")
            ftitle = finding.get("title", finding.get("description", ""))
            detail = finding.get("detail", finding.get("description", ""))
            rec = finding.get("recommendation", "")
            add_finding_box(doc, sev, ftitle, detail, rec)

        # Try to add relevant image
        domain_lower = domain.lower().replace(" ", "_")
        for img_key in (f"nets_{domain_lower}", domain_lower, f"board_{domain_lower}"):
            if img_key in img:
                figure_num += 1
                add_image_with_caption(doc, img[img_key],
                    f"Figure {figure_num}: {domain} — highlighted traces and components.",
                    width_inches=5.5)
                break

        doc.add_page_break()

    # ---- Net renders ----
    net_imgs = {k: v for k, v in img.items() if k.startswith("net_")}
    if net_imgs:
        doc.add_heading("Net Highlight Renders", level=1)
        doc.add_paragraph(
            "The following net renders show individual signal paths highlighted on the board layout. "
            "These identify routing, via transitions, and connection topology for key signals."
        )
        for label, path in sorted(net_imgs.items()):
            net_name = label.replace("net_", "").replace("_", " ").upper()
            figure_num += 1
            add_image_with_caption(doc, path,
                f"Figure {figure_num}: Net '{net_name}' — traces and vias highlighted.",
                width_inches=5.0)

    # ---- Drill table ----
    if design.vias:
        doc.add_heading("Drill Table & Via Analysis", level=1)
        drill_counts: dict[float, int] = {}
        for v in design.vias:
            d = v.get("drill_mm", 0) if isinstance(v, dict) else getattr(v, "drill_mm", 0)
            if d > 0:
                d_rounded = round(d, 3)
                drill_counts[d_rounded] = drill_counts.get(d_rounded, 0) + 1
        if drill_counts:
            drill_rows = [(f"{sz:.3f}", str(cnt)) for sz, cnt in sorted(drill_counts.items())]
            drill_rows.append(("Total", str(sum(drill_counts.values()))))
            add_styled_table(doc, ["Drill Size (mm)", "Count"], drill_rows,
                            col_widths=[3.0, 3.5])

    # ---- Footer ----
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Report —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated by MCP PCB EMCopilot | {author}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    return os.path.abspath(output_path)
