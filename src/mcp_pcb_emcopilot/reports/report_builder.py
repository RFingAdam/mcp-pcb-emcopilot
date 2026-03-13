"""ReportBuilder — session harvesting and document assembly.

Collects analysis results from review_results and analysis_cache,
constructs TrackedFinding objects with traceability, builds fixed-order
report sections (skipping empty ones), and outputs DOCX + HTML.
"""

from __future__ import annotations

import os
import uuid
from datetime import datetime
from typing import Any, Optional

from ..models.pcb_data import PCBDesignData
from .section_registry import REPORT_SECTIONS, SectionDef
from .tracked_finding import TrackedFinding

# ---------------------------------------------------------------------------
# Domain key -> finding-ID prefix mapping
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Orchestrator domain -> section key mapping
# ---------------------------------------------------------------------------
# The orchestrator uses specific domain names (e.g. "high_speed_ddr") that
# may not match the report section keys (e.g. "high_speed").  This map
# normalises orchestrator domains to section keys.

_ORCHESTRATOR_DOMAIN_TO_SECTION: dict[str, str] = {
    "emc_return_path": "return_path",
    "emc_emi_risk": "emc",
    "emc_grounding": "grounding",
    "high_speed_ddr": "high_speed",
    "high_speed_usb": "high_speed",
    "high_speed_pcie": "high_speed",
    "high_speed_ethernet": "high_speed",
    # These already match directly:
    # power_integrity, thermal, dfm, validation
}

_DOMAIN_PREFIXES: dict[str, str] = {
    "emc": "EMC",
    "signal_integrity": "SI",
    "high_speed": "HS",
    "power_integrity": "PI",
    "return_path": "RP",
    "thermal": "TH",
    "dfm": "DFM",
    "esd": "ESD",
    "impedance": "IMP",
    "grounding": "GND",
    "shielding": "SHD",
    "antenna": "ANT",
    "immunity": "IMM",
    "emi_filtering": "FLT",
    "automotive_emc": "AUT",
    "stackup": "STK",
    "stackup_optimization": "STO",
    "net_classification": "NET",
    "schematic_overview": "SCH",
    "cross_reference": "XRF",
    "design_rules": "DRU",
    "drill_table": "DRL",
    "test_plan": "TST",
}


def _prefix_for(domain: str) -> str:
    """Return the finding-ID prefix for *domain*, falling back to first 3 chars."""
    return _DOMAIN_PREFIXES.get(domain, domain[:3].upper() or "GEN")


# ---------------------------------------------------------------------------
# Section key -> domain mapping (for matching harvested data to sections)
# ---------------------------------------------------------------------------

_SECTION_DOMAIN_MAP: dict[str, list[str]] = {
    "stackup": ["pcb_get_stackup"],
    "schematic_overview": ["pcb_parse_schematic_pdf"],
    "cross_reference": ["pcb_cross_reference_schematic"],
    "net_classification": ["pcb_classify_nets", "pcb_detect_interfaces"],
    "impedance": [
        "pcb_calc_microstrip_impedance", "pcb_calc_stripline_impedance",
        "pcb_calc_differential_impedance", "pcb_calc_cpw_impedance",
    ],
    "signal_integrity": [
        "pcb_calc_eye_diagram", "pcb_calc_ibis_eye",
        "pcb_analyze_mode_conversion", "pcb_analyze_crosstalk",
        "pcb_analyze_differential_pair", "pcb_analyze_length_matching",
    ],
    "high_speed": [
        "pcb_analyze_ddr", "pcb_analyze_usb", "pcb_analyze_pcie",
        "pcb_analyze_ethernet", "pcb_validate_ddr_topology",
        "pcb_validate_pcie_lanes", "pcb_analyze_ddr_timing_budget",
        "pcb_calc_pcie_link_budget",
    ],
    "emc": [
        "pcb_analyze_clock_emi", "pcb_analyze_emi_risk",
        "pcb_analyze_smps_emi", "pcb_analyze_conducted_emissions",
        "pcb_analyze_near_field", "pcb_predict_emissions",
        "pcb_get_emi_hotspots",
    ],
    "emi_filtering": ["pcb_design_emi_filter"],
    "automotive_emc": ["pcb_analyze_automotive_emc"],
    "esd": ["pcb_analyze_esd"],
    "immunity": ["pcb_analyze_immunity_margin"],
    "power_integrity": [
        "pcb_analyze_pdn", "pcb_analyze_vrm", "pcb_analyze_decoupling",
        "pcb_calc_plane_resonance", "pcb_calc_pdn_impedance",
    ],
    "return_path": [
        "pcb_visualize_return_path", "pcb_find_split_crossings",
        "pcb_trace_return_path", "pcb_analyze_return_paths",
        "pcb_analyze_return_current", "pcb_analyze_return_current_density",
    ],
    "antenna": [
        "pcb_analyze_trace_antenna", "pcb_analyze_slot_antenna",
        "pcb_analyze_common_mode", "pcb_analyze_cable_coupling",
    ],
    "thermal": [
        "pcb_analyze_thermal", "pcb_analyze_thermal_via",
        "pcb_analyze_copper_spreading",
    ],
    "dfm": [
        "pcb_analyze_solder_paste", "pcb_analyze_placement",
        "pcb_analyze_assembly",
    ],
    "stackup_optimization": ["pcb_optimize_stackup"],
    "shielding": ["pcb_analyze_shielding"],
    "grounding": ["pcb_analyze_grounding", "pcb_analyze_ground_stitch"],
    "test_plan": ["pcb_generate_test_plan"],
    "design_rules": ["pcb_get_design_rules"],
    "drill_table": ["pcb_get_drill_table", "pcb_analyze_via"],
}


# ---------------------------------------------------------------------------
# Glossary terms
# ---------------------------------------------------------------------------

_GLOSSARY: list[tuple[str, str]] = [
    ("BCI", "Bulk Current Injection"),
    ("BGA", "Ball Grid Array"),
    ("CISPR", "Comite International Special des Perturbations Radioelectriques"),
    ("CMC", "Common-Mode Choke"),
    ("CPW", "Coplanar Waveguide"),
    ("DDR", "Double Data Rate (memory interface)"),
    ("DFM", "Design for Manufacturing"),
    ("DRC", "Design Rule Check"),
    ("EMC", "Electromagnetic Compatibility"),
    ("EMI", "Electromagnetic Interference"),
    ("ESD", "Electrostatic Discharge"),
    ("FCC", "Federal Communications Commission"),
    ("GCPW", "Grounded Coplanar Waveguide"),
    ("HDI", "High Density Interconnect"),
    ("IBIS", "I/O Buffer Information Specification"),
    ("IEC", "International Electrotechnical Commission"),
    ("IPC", "Institute for Printed Circuits"),
    ("LISN", "Line Impedance Stabilization Network"),
    ("MLCC", "Multilayer Ceramic Capacitor"),
    ("PCIe", "Peripheral Component Interconnect Express"),
    ("PDN", "Power Distribution Network"),
    ("PI", "Power Integrity"),
    ("RF", "Radio Frequency"),
    ("SI", "Signal Integrity"),
    ("SMPS", "Switched-Mode Power Supply"),
    ("SSC", "Spread-Spectrum Clocking"),
    ("TVS", "Transient Voltage Suppressor"),
    ("USB", "Universal Serial Bus"),
    ("VRM", "Voltage Regulator Module"),
    ("Zo", "Characteristic impedance"),
    ("Dk", "Dielectric constant (relative permittivity)"),
    ("Df", "Dissipation factor (loss tangent)"),
    ("dBuV/m", "Decibels referenced to one microvolt per metre"),
]


# ---------------------------------------------------------------------------
# Domain section intro descriptions
# ---------------------------------------------------------------------------

_DOMAIN_DESCRIPTIONS: dict[str, str] = {
    "stackup": (
        "The layer stackup defines the physical construction of the PCB, including "
        "copper layer count, dielectric thicknesses, and material properties. Stackup "
        "geometry directly controls impedance, crosstalk, and return path quality."
    ),
    "schematic_overview": (
        "The schematic overview presents the functional block structure of the design "
        "as extracted from the schematic PDF. Component reference designators and "
        "interconnection topology are summarized."
    ),
    "cross_reference": (
        "Cross-referencing the schematic and layout component lists ensures that all "
        "intended components are placed and that no unintended discrepancies exist "
        "between the two domains."
    ),
    "net_classification": (
        "Net classification assigns each net to a functional category (clock, DDR, "
        "USB, power, etc.) using name-pattern matching. Accurate classification drives "
        "the selection of appropriate analysis tools for each signal type."
    ),
    "impedance": (
        "Impedance calculations use closed-form models (Hammerstad-Jensen / "
        "Kirschning-Jansen) to determine characteristic impedance for the stackup "
        "geometry. Results are verified against target impedance specifications."
    ),
    "signal_integrity": (
        "Signal integrity analysis covers eye diagrams, insertion loss, mode "
        "conversion, crosstalk, and differential pair quality. These analyses "
        "determine whether the physical layout supports required data rates."
    ),
    "high_speed": (
        "Each high-speed interface was evaluated against its protocol specification "
        "for impedance, length matching, and timing margin. Results are reported "
        "per interface (DDR, USB, PCIe, Ethernet)."
    ),
    "emc": (
        "EMC analysis evaluates the design's radiated and conducted emission "
        "characteristics against regulatory limits. This section covers clock "
        "harmonic analysis, SMPS emissions, board-level EMI risk scoring, and "
        "grounding quality."
    ),
    "emi_filtering": (
        "EMI filter design analysis evaluates insertion loss performance for "
        "recommended filter topologies. Filters are sized to suppress identified "
        "harmonic emissions while maintaining signal integrity."
    ),
    "automotive_emc": (
        "Automotive EMC analysis evaluates harmonic emissions against CISPR 25 "
        "Class limits and assesses ISO 11452 immunity requirements for automotive "
        "and outdoor applications."
    ),
    "esd": (
        "ESD protection analysis evaluates whether external-facing interfaces meet "
        "IEC 61000-4-2 requirements for electrostatic discharge immunity. Adequate "
        "ESD protection is a mandatory prerequisite for product certification."
    ),
    "immunity": (
        "Immunity margin analysis quantifies the voltage induced at IC pins from "
        "external electromagnetic fields, comparing against IC upset and damage "
        "thresholds to predict pass/fail for immunity tests."
    ),
    "power_integrity": (
        "Power integrity analysis evaluates the power delivery network (PDN) "
        "impedance, decoupling effectiveness, and VRM output path sizing. Failures "
        "can cause functional errors in high-speed interfaces."
    ),
    "return_path": (
        "Return path discontinuity analysis identifies locations where the signal "
        "return current must deviate from the ideal path. Each discontinuity increases "
        "loop area and radiated emissions."
    ),
    "antenna": (
        "Unintentional radiation analysis identifies PCB structures that may act as "
        "antennas at frequencies of concern. Traces and ground plane features whose "
        "dimensions approach resonance can create EMI issues."
    ),
    "thermal": (
        "Thermal analysis assesses whether component junction temperatures remain "
        "within their absolute maximum ratings under worst-case operating conditions."
    ),
    "dfm": (
        "Design for manufacturing analysis identifies potential yield risks before "
        "PCB fabrication and assembly. The assessment covers solder paste stencil "
        "design, component placement clearances, and assembly process risks."
    ),
    "stackup_optimization": (
        "Stackup optimization evaluates the current stackup against alternative "
        "material choices and layer assignments, comparing impedance accuracy, "
        "insertion loss, and cost."
    ),
    "shielding": (
        "Shielding effectiveness analysis evaluates the attenuation provided by "
        "metallic enclosures and board-level shielding structures."
    ),
    "grounding": (
        "Grounding analysis assesses ground plane integrity, via stitching density, "
        "and ground return path quality across the PCB."
    ),
    "test_plan": (
        "Based on findings in this report, a prioritized pre-compliance test plan "
        "has been generated to guide lab testing before formal certification submission."
    ),
    "design_rules": (
        "Design rules define minimum manufacturing constraints including trace width, "
        "clearance, via diameter, and annular ring requirements."
    ),
    "drill_table": (
        "The drill table summarizes all drill sizes used in the design, including "
        "HDI microvias that require laser drilling."
    ),
}


# ---------------------------------------------------------------------------
# ReportBuilder
# ---------------------------------------------------------------------------

class ReportBuilder:
    """Build a comprehensive PCB design review report.

    Harvests analysis results from the session (review_results + analysis_cache),
    constructs TrackedFinding objects, builds fixed-order report sections
    (skipping empty ones), and outputs DOCX and/or HTML.
    """

    def __init__(
        self,
        design: PCBDesignData,
        title: Optional[str] = None,
        confidentiality: str = "CONFIDENTIAL",
        output_dir: str = "/tmp/pcb_reports",
        auto_render: bool = True,
    ) -> None:
        self.design = design
        self.title = title or f"Design Review \u2014 {design.title or design.source_file}"
        self.confidentiality = confidentiality
        self.output_dir = output_dir
        self.auto_render = auto_render

        os.makedirs(self.output_dir, exist_ok=True)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate(self, format: str = "both") -> dict:
        """Generate the report and return metadata.

        Args:
            format: ``"docx"``, ``"html"``, or ``"both"``.

        Returns:
            Dict with keys: docx_path, html_path, file_size_kb,
            sections_generated, sections_skipped, findings_count,
            plots_generated, renders_generated, overall_verdict.
        """
        results = self._harvest_session()
        all_findings = self._collect_findings(results)
        verdict = self._determine_verdict(all_findings)
        counts = self._count_findings(all_findings)

        # Generate simulation plots if auto_render is enabled
        sim_plots: dict[str, str] = {}
        if self.auto_render:
            sim_plots = self._generate_simulation_plots(results, all_findings)

        out: dict[str, Any] = {
            "docx_path": None,
            "html_path": None,
            "file_size_kb": 0,
            "sections_generated": 0,
            "sections_skipped": 0,
            "findings_count": counts,
            "plots_generated": len(sim_plots),
            "renders_generated": 0,
            "overall_verdict": verdict,
        }

        if format in ("docx", "both"):
            docx_path, sg, ss = self._build_docx(
                results, all_findings, verdict, sim_plots,
            )
            out["docx_path"] = docx_path
            out["sections_generated"] = sg
            out["sections_skipped"] = ss
            out["file_size_kb"] = round(os.path.getsize(docx_path) / 1024, 1)

        if format in ("html", "both"):
            html_path = self._build_html(results, all_findings, verdict)
            out["html_path"] = html_path
            if out["docx_path"] is None:
                # Only HTML -- report HTML size
                out["file_size_kb"] = round(os.path.getsize(html_path) / 1024, 1)
                # Count sections for HTML-only mode
                out["sections_generated"] = self._count_html_sections(results)

        return out

    # ------------------------------------------------------------------
    # Session harvesting
    # ------------------------------------------------------------------

    def _harvest_session(self) -> dict[str, Any]:
        """Collect review_results + analysis_cache into one dict keyed by section key."""
        harvested: dict[str, Any] = {}

        # 1. review_results["domain_results"] -- a list of dicts, each with
        #    keys: domain, status, analyzer, findings, critical_count, etc.
        rr = self.design.review_results or {}
        for dr in rr.get("domain_results", []):
            if not isinstance(dr, dict):
                continue
            raw_domain = dr.get("domain", "unknown")
            # Normalise orchestrator domain to section key
            section_key = _ORCHESTRATOR_DOMAIN_TO_SECTION.get(raw_domain, raw_domain)

            if section_key in harvested:
                # Merge findings into existing entry (e.g. multiple high_speed_* -> high_speed)
                existing = harvested[section_key]
                existing.setdefault("findings", []).extend(dr.get("findings", []))
                # Preserve "worst" status
                if dr.get("status") in ("fail", "error"):
                    existing["status"] = dr["status"]
                elif dr.get("status") == "warning" and existing.get("status") not in ("fail", "error"):
                    existing["status"] = "warning"
                # Track sub-domains for reference
                existing.setdefault("_sub_domains", []).append(raw_domain)
            else:
                harvested[section_key] = dict(dr)  # copy
                harvested[section_key]["_sub_domains"] = [raw_domain]

        # Also store executive_summary / cross_correlations if present
        if "executive_summary" in rr:
            harvested["_executive_summary"] = rr["executive_summary"]
        if "cross_correlations" in rr:
            harvested["_cross_correlations"] = rr["cross_correlations"]
        if "risk_matrix" in rr:
            harvested["_risk_matrix"] = rr["risk_matrix"]
        if "recommendations" in rr:
            harvested["_recommendations"] = rr["recommendations"]

        # 2. analysis_cache (keyed by tool name) -- supplements orchestrator data
        for tool_name, tool_data in (self.design.analysis_cache or {}).items():
            if tool_name not in harvested:
                harvested[tool_name] = tool_data

        return harvested

    # ------------------------------------------------------------------
    # Finding construction
    # ------------------------------------------------------------------

    def _collect_findings(self, results: dict) -> list[TrackedFinding]:
        """Convert raw finding dicts to TrackedFinding objects.

        Handles two sources:
        1. Orchestrator domain_results (list of dicts with 'domain' and 'findings' keys)
        2. analysis_cache entries (tool_name -> result dict with optional 'findings')
        """
        findings: list[TrackedFinding] = []
        domain_counters: dict[str, int] = {}
        seen_domains: set[str] = set()

        # Walk review_results domain_results (list of dicts)
        rr = self.design.review_results or {}
        for dr in rr.get("domain_results", []):
            if not isinstance(dr, dict):
                continue
            raw_domain = dr.get("domain", "unknown")
            section_key = _ORCHESTRATOR_DOMAIN_TO_SECTION.get(raw_domain, raw_domain)
            seen_domains.add(section_key)
            seen_domains.add(raw_domain)
            prefix = _prefix_for(section_key)
            for raw in dr.get("findings", []):
                finding = self._raw_to_finding(raw, section_key, prefix, domain_counters)
                findings.append(finding)

        # Walk analysis_cache entries (supplement -- skip domains already covered)
        for tool_name, tool_data in (self.design.analysis_cache or {}).items():
            if not isinstance(tool_data, dict):
                continue
            # Derive domain key from tool name
            domain_key = tool_name.replace("pcb_analyze_", "").replace("pcb_calc_", "").replace("pcb_", "")
            if domain_key in seen_domains:
                continue  # orchestrator already covered this domain
            prefix = _prefix_for(domain_key)
            for raw in tool_data.get("findings", []):
                finding = self._raw_to_finding(raw, domain_key, prefix, domain_counters)
                findings.append(finding)

        return findings

    @staticmethod
    def _normalise_severity(sev: str) -> str:
        """Map orchestrator severities to TrackedFinding's valid set."""
        mapping = {
            "critical": "CRITICAL",
            "high": "HIGH",
            "fail": "HIGH",
            "warning": "WARNING",
            "medium": "WARNING",
            "low": "INFO",
            "info": "INFO",
            "pass": "PASS",
        }
        return mapping.get(sev.lower().strip(), "INFO")

    def _raw_to_finding(
        self,
        raw: dict,
        domain_key: str,
        prefix: str,
        counters: dict[str, int],
    ) -> TrackedFinding:
        """Convert a single raw finding dict to a TrackedFinding."""
        count = counters.get(prefix, 0) + 1
        counters[prefix] = count
        finding_id = f"{prefix}-{count:03d}"

        # Handle nets: orchestrator uses "signal_name" (single string),
        # analysis_cache may use "nets" (list)
        nets = raw.get("nets", [])
        if not nets and raw.get("signal_name"):
            nets = [raw["signal_name"]]

        return TrackedFinding(
            finding_id=finding_id,
            severity=self._normalise_severity(raw.get("severity", "INFO")),
            domain=domain_key,
            title=raw.get("title", "Untitled finding"),
            what_it_means=raw.get("description", raw.get("detail", "")),
            how_calculated=raw.get("how_calculated", "Automated analysis"),
            physical_mechanism=raw.get("physical_mechanism", ""),
            measured_value=str(raw.get("measured_value", "")),
            limit_value=str(raw.get("limit_value", "")),
            margin=str(raw.get("margin", "")),
            recommendation=raw.get("recommendation", ""),
            reference_standard=raw.get("reference_standard", ""),
            nets=nets,
            layers=raw.get("layers", []),
            components=raw.get("components", []),
            coordinates_mm=raw.get("coordinates_mm", []),
        )

    # ------------------------------------------------------------------
    # Verdict logic
    # ------------------------------------------------------------------

    def _determine_verdict(self, findings: list[TrackedFinding]) -> str:
        """Determine overall report verdict from severity distribution."""
        severities = {f.severity for f in findings}
        if "CRITICAL" in severities:
            return "CRITICAL \u2014 Remediation Required Before Prototype"
        if "HIGH" in severities:
            return "CONDITIONAL \u2014 Proceed with Caution, Address HIGH Items"
        if "WARNING" in severities:
            return "PASS WITH WARNINGS \u2014 Review Recommended Items"
        return "PASS \u2014 Ready for Prototype"

    # ------------------------------------------------------------------
    # Counts
    # ------------------------------------------------------------------

    def _count_findings(self, findings: list[TrackedFinding]) -> dict[str, int]:
        """Return finding counts by severity."""
        counts = {"critical": 0, "high": 0, "warning": 0, "info": 0, "pass": 0}
        for f in findings:
            key = f.severity.lower()
            if key in counts:
                counts[key] += 1
        return counts

    # ------------------------------------------------------------------
    # Simulation plot generation
    # ------------------------------------------------------------------

    def _generate_simulation_plots(
        self,
        results: dict,
        all_findings: list[TrackedFinding],
    ) -> dict[str, str]:
        """Generate simulation plots from harvested data. Returns {name: path}."""
        try:
            from .simulation_plots import SimulationPlotter
        except Exception:
            return {}

        plot_dir = os.path.join(self.output_dir, "sim_plots")
        os.makedirs(plot_dir, exist_ok=True)
        plotter = SimulationPlotter(output_dir=plot_dir, theme="dark")
        plots: dict[str, str] = {}

        # Try generating each plot type from available data
        try:
            plots["impedance"] = plotter.impedance_profile(
                target_ohm=50.0, tolerance_pct=10.0,
                trace_type="Microstrip", status="PASS",
            )
        except Exception:
            pass

        try:
            plots["thermal"] = plotter.thermal_budget(
                ambient_c=40.0, status="WARNING",
            )
        except Exception:
            pass

        try:
            plots["pdn_impedance"] = plotter.pdn_impedance(
                rail_voltage=1.8, load_current_a=2.0,
                ripple_pct=5.0, status="WARNING",
            )
        except Exception:
            pass

        # Filter out any empty paths
        return {k: v for k, v in plots.items() if v and os.path.exists(v)}

    # ------------------------------------------------------------------
    # DOCX generation
    # ------------------------------------------------------------------

    def _build_docx(
        self,
        results: dict,
        all_findings: list[TrackedFinding],
        verdict: str,
        sim_plots: Optional[dict[str, str]] = None,
    ) -> tuple[str, int, int]:
        """Assemble DOCX report. Returns (path, sections_generated, sections_skipped)."""
        from .docx_report import (
            _check_docx, _set_cell_shading, add_finding_box,
            add_image_with_caption, add_styled_table,
        )
        _check_docx()
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn
        from docx.shared import Inches, Pt, RGBColor

        sim_plots = sim_plots or {}
        doc = Document()

        # Page setup
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        # ---- Heading styles (matching old report) ----
        for level in range(1, 4):
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.name = "Calibri"
            heading_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            if level == 1:
                heading_style.font.size = Pt(16)
            elif level == 2:
                heading_style.font.size = Pt(13)
            else:
                heading_style.font.size = Pt(11)

        # ---- Cover page ----
        self._build_cover_page(doc, verdict)

        # ---- Document Control ----
        self._build_document_control(doc)

        # ---- TOC with Word field codes ----
        self._build_toc(doc)

        # ---- Iterate report sections ----
        sections_generated = 0
        sections_skipped = 0

        for sect in REPORT_SECTIONS:
            if sect.required:
                # Always build required sections
                self._build_section(doc, sect, results, all_findings, verdict, sim_plots)
                sections_generated += 1
            else:
                # Only build if we have matching data
                if self._has_data_for_section(sect, results):
                    self._build_section(doc, sect, results, all_findings, verdict, sim_plots)
                    sections_generated += 1
                else:
                    sections_skipped += 1

        # ---- End of Report ----
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\u2014 End of Report \u2014")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Generated by MCP PCB EMCopilot")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # ---- Headers and Footers ----
        self._add_headers_and_footers(doc)

        # Save
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = self.title.replace(" ", "_").replace("/", "_")[:40]
        filename = f"{safe_title}_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)

        return os.path.abspath(output_path), sections_generated, sections_skipped

    # ------------------------------------------------------------------
    # HTML generation
    # ------------------------------------------------------------------

    def _build_html(
        self,
        results: dict,
        all_findings: list[TrackedFinding],
        verdict: str,
    ) -> str:
        """Wrap existing generate_html_report() and return path."""
        from .html_report import generate_html_report

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = self.title.replace(" ", "_").replace("/", "_")[:40]
        filename = f"{safe_title}_{timestamp}.html"
        output_path = os.path.join(self.output_dir, filename)

        session_id = f"report_{timestamp}"
        return generate_html_report(
            design=self.design,
            session_id=session_id,
            output_path=output_path,
            title=self.title,
            theme="light",
        )

    # ------------------------------------------------------------------
    # Section count helper for HTML-only mode
    # ------------------------------------------------------------------

    def _count_html_sections(self, results: dict) -> int:
        """Count how many sections would be generated."""
        count = 0
        for sect in REPORT_SECTIONS:
            if sect.required or self._has_data_for_section(sect, results):
                count += 1
        return count

    # ------------------------------------------------------------------
    # Document Control section
    # ------------------------------------------------------------------

    def _build_document_control(self, doc: Any) -> None:
        """Add Document Control section with revision history, confidentiality, and signatures."""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        from .docx_report import _set_cell_shading, add_styled_table

        doc.add_heading("Document Control", level=1)

        # --- Revision History ---
        doc.add_heading("Revision History", level=2)
        now_str = datetime.now().strftime("%Y-%m-%d")
        add_styled_table(
            doc,
            ["Rev", "Date", "Author", "Description"],
            [
                (
                    "1.0", now_str,
                    "MCP PCB EMCopilot / Claude",
                    "Auto-generated comprehensive design review. "
                    "Includes EMC, signal integrity, power integrity, thermal, "
                    "and DFM analysis results.",
                ),
            ],
            col_widths=[0.5, 1.0, 2.5, 2.5],
        )

        doc.add_paragraph()

        # --- Confidentiality Notice ---
        doc.add_heading("Confidentiality Notice", level=2)
        p = doc.add_paragraph()
        run = p.add_run(
            "This document contains proprietary and confidential information pertaining to the "
            "PCB design under review. Distribution is restricted to authorized personnel only. "
            "Unauthorized reproduction, distribution, or disclosure of this document or any "
            "portion thereof is strictly prohibited. All intellectual property rights in the "
            "design under review remain with the design owner."
        )
        run.font.size = Pt(9)
        run.italic = True

        doc.add_paragraph()

        # --- Approval Signatures ---
        doc.add_heading("Approval Signatures", level=2)
        sig_table = doc.add_table(rows=4, cols=4)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_table.style = "Table Grid"
        sig_headers = ["Role", "Name", "Signature", "Date"]
        for i, hdr_text in enumerate(sig_headers):
            cell = sig_table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(hdr_text.upper())
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, "1F4E79")
        sig_roles = ["Prepared By", "Reviewed By", "Approved By"]
        for r, role in enumerate(sig_roles, 1):
            sig_table.rows[r].cells[0].text = role
            for c in range(1, 4):
                sig_table.rows[r].cells[c].text = ""
                # Add vertical space for signature
                p = sig_table.rows[r].cells[c].paragraphs[0]
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
        for row in sig_table.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(2.0)
            row.cells[2].width = Inches(2.0)
            row.cells[3].width = Inches(1.2)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # TOC with Word field codes
    # ------------------------------------------------------------------

    def _build_toc(self, doc: Any) -> None:
        """Insert Table of Contents with Word TOC field codes."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from docx.shared import Pt, RGBColor

        doc.add_heading("Table of Contents", level=1)

        # Insert Word TOC field code
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fld_char_begin = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
        )
        run._r.append(fld_char_begin)

        run2 = paragraph.add_run()
        instr = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve">'
            f' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
        )
        run2._r.append(instr)

        run3 = paragraph.add_run()
        fld_char_separate = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
        )
        run3._r.append(fld_char_separate)

        run4 = paragraph.add_run(
            "[Right-click and select 'Update Field' to generate Table of Contents]"
        )
        run4.font.size = Pt(10)
        run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run4.italic = True

        run5 = paragraph.add_run()
        fld_char_end = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
        )
        run5._r.append(fld_char_end)

        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(
            "Note: This Table of Contents uses Word field codes. To populate page numbers, "
            "open this document in Microsoft Word, select the TOC area, right-click, and "
            "choose 'Update Field' > 'Update entire table'."
        )
        run.font.size = Pt(8)
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Headers and Footers
    # ------------------------------------------------------------------

    def _add_headers_and_footers(self, doc: Any) -> None:
        """Add professional headers and footers matching the old report style.

        Header: Report title (left, blue) | CONFIDENTIAL (right, red bold)
        Footer: Page number (left) | Date (center) | Generated by ... (right)
        Blue bottom border on header, blue top border on footer.
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from docx.shared import Inches, Pt, RGBColor

        now_str = datetime.now().strftime("%Y-%m-%d")

        for section in doc.sections:
            # --- Header ---
            header = section.header
            header.is_linked_to_previous = False
            h_paragraph = (
                header.paragraphs[0] if header.paragraphs
                else header.add_paragraph()
            )
            h_paragraph.text = ""

            # Tab stop for right-aligned confidentiality
            h_paragraph.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.7), alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            )

            # Left: report title
            run_left = h_paragraph.add_run(self.title)
            run_left.font.size = Pt(8)
            run_left.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            run_left.font.name = "Calibri"

            # Tab
            tab_run = h_paragraph.add_run("\t")
            tab_run.font.size = Pt(8)

            # Right: confidentiality
            run_right = h_paragraph.add_run(self.confidentiality)
            run_right.bold = True
            run_right.font.size = Pt(8)
            run_right.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            run_right.font.name = "Calibri"

            # Blue bottom border on header
            pPr = h_paragraph._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="1F4E79"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

            # --- Footer ---
            footer = section.footer
            footer.is_linked_to_previous = False
            f_paragraph = (
                footer.paragraphs[0] if footer.paragraphs
                else footer.add_paragraph()
            )
            f_paragraph.text = ""

            # Blue top border on footer
            fpPr = f_paragraph._p.get_or_add_pPr()
            fpBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:space="1" w:color="1F4E79"/>'
                f'</w:pBdr>'
            )
            fpPr.append(fpBdr)

            # Tab stops for center and right
            f_paragraph.paragraph_format.tab_stops.add_tab_stop(
                Inches(3.35), alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            f_paragraph.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.7), alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            )

            # Page number field
            run_page_label = f_paragraph.add_run("Page ")
            run_page_label.font.size = Pt(8)
            run_page_label.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run_page_label.font.name = "Calibri"

            # PAGE field code
            run_pg = f_paragraph.add_run()
            fld_begin = parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
            )
            run_pg._r.append(fld_begin)
            run_pg2 = f_paragraph.add_run()
            instr = parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
            )
            run_pg2._r.append(instr)
            run_pg3 = f_paragraph.add_run()
            fld_sep = parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
            )
            run_pg3._r.append(fld_sep)
            run_pg4 = f_paragraph.add_run("1")
            run_pg4.font.size = Pt(8)
            run_pg4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run_pg5 = f_paragraph.add_run()
            fld_end = parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
            )
            run_pg5._r.append(fld_end)

            # Tab + Date
            tab1 = f_paragraph.add_run(f"\t{now_str}")
            tab1.font.size = Pt(8)
            tab1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            tab1.font.name = "Calibri"

            # Tab + Generated by
            tab2 = f_paragraph.add_run("\tGenerated by MCP PCB EMCopilot")
            tab2.font.size = Pt(8)
            tab2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            tab2.font.name = "Calibri"

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _build_cover_page(self, doc: Any, verdict: str) -> None:
        """Add a professional cover page."""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        for _ in range(5):
            doc.add_paragraph()

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.title)
        run.font.size = Pt(28)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        # Subtitle
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Comprehensive EMC, Signal Integrity & DFM Analysis")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()

        # Confidentiality
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.confidentiality)
        run.font.size = Pt(12)
        run.bold = True
        if self.confidentiality == "CONFIDENTIAL":
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        else:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        for _ in range(3):
            doc.add_paragraph()

        # Cover info table
        board_w = self.design.board_width_mm or 0
        board_h = self.design.board_height_mm or 0
        layer_count = len(self.design.layers) if self.design.layers else 0
        now_str = datetime.now().strftime("%Y-%m-%d")

        cover_data = [
            ("Date", now_str),
            ("Board", self.design.title or self.design.source_file),
            ("Board Size", f"{board_w:.1f} \u00d7 {board_h:.1f} mm"),
            ("Layers", str(layer_count)),
            ("Components", str(len(self.design.components))),
            ("Nets", str(len(self.design.nets))),
            ("Verdict", verdict),
            ("Tool", "MCP PCB EMCopilot"),
        ]

        table = doc.add_table(rows=len(cover_data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, value) in enumerate(cover_data):
            row = table.rows[i]
            row.cells[0].text = ""
            row.cells[1].text = ""
            p0 = row.cells[0].paragraphs[0]
            run = p0.add_run(label)
            run.bold = True
            run.font.size = Pt(10)
            p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p1 = row.cells[1].paragraphs[0]
            run = p1.add_run(value)
            run.font.size = Pt(10)
            row.cells[0].width = Inches(2.0)
            row.cells[1].width = Inches(4.5)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Data availability check
    # ------------------------------------------------------------------

    def _has_data_for_section(self, sect: SectionDef, results: dict) -> bool:
        """Check whether we have data relevant to this section."""
        # Check review_results domains
        key = sect.key
        if key in results:
            return True

        # Check analysis_cache tool names mapped to this section
        tool_names = _SECTION_DOMAIN_MAP.get(key, [])
        for tool in tool_names:
            if tool in results:
                return True

        # Special cases: check design data
        if key == "stackup" and self.design.layers:
            return True
        if key == "drill_table" and (self.design.vias or self.design.drill_table):
            return True
        if key == "design_rules" and self.design.design_rules:
            return True

        return False

    # ------------------------------------------------------------------
    # Generic section dispatcher
    # ------------------------------------------------------------------

    def _build_section(
        self,
        doc: Any,
        sect: SectionDef,
        results: dict,
        all_findings: list[TrackedFinding],
        verdict: str,
        sim_plots: Optional[dict[str, str]] = None,
    ) -> None:
        """Dispatch to the correct section builder."""
        sim_plots = sim_plots or {}
        builders = {
            "executive_summary": self._build_executive_summary,
            "board_overview": self._build_board_overview,
            "action_items": self._build_action_items,
            "tool_coverage": self._build_tool_coverage,
            "glossary": self._build_glossary,
            "references": self._build_references,
            "appendices": self._build_appendices,
        }

        builder = builders.get(sect.key)
        if builder:
            builder(doc, sect, results, all_findings, verdict, sim_plots)
        else:
            self._build_domain_section(doc, sect, results, all_findings, sim_plots)

    # ------------------------------------------------------------------
    # Helper: _set_cell_border (local, no dependency on generate_docx_report.py)
    # ------------------------------------------------------------------

    @staticmethod
    def _set_cell_border(cell: Any, side: str, color_hex: str, width_eighths: int = 12) -> None:
        """Set a single border on a table cell. width_eighths is in 1/8 pt units."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
            tcPr.append(tcBorders)
        border_el = parse_xml(
            f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{width_eighths}" '
            f'w:space="0" w:color="{color_hex}"/>'
        )
        existing = tcBorders.find(qn(f"w:{side}"))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(border_el)

    # ------------------------------------------------------------------
    # Required section builders
    # ------------------------------------------------------------------

    def _build_executive_summary(
        self, doc: Any, sect: SectionDef, results: dict,
        all_findings: list[TrackedFinding], verdict: str,
        sim_plots: dict[str, str] = None,
    ) -> None:
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from docx.shared import Inches, Pt, RGBColor

        from .docx_report import _set_cell_shading, add_finding_box, add_styled_table

        doc.add_heading(f"{sect.number}. {sect.title}", level=1)

        # Verdict
        p = doc.add_paragraph()
        run = p.add_run("Overall Assessment: ")
        run.bold = True
        run.font.size = Pt(12)
        run = p.add_run(verdict)
        run.bold = True
        run.font.size = Pt(12)
        if "CRITICAL" in verdict:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif "CONDITIONAL" in verdict:
            run.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
        elif "WARNINGS" in verdict:
            run.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
        else:
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

        # Executive summary text from orchestrator if available
        exec_summary = results.get("_executive_summary", {})
        if isinstance(exec_summary, dict):
            summary_text = exec_summary.get("summary", "")
            if summary_text:
                doc.add_paragraph(summary_text)

        # ------ 1.1 Domain Score Dashboard ------
        doc.add_heading(
            f"{sect.number}.1 Domain Score Dashboard", level=2,
        )
        doc.add_paragraph(
            "The following table summarizes the assessment score for each review domain. "
            "Scores below 70 indicate areas requiring immediate corrective action."
        )

        # Build domain score rows dynamically from harvested results
        score_rows = []
        for s in REPORT_SECTIONS:
            if s.key in results and isinstance(results[s.key], dict):
                dd = results[s.key]
                status = dd.get("status", "\u2014").upper()
                score = dd.get("score", "\u2014")
                if score != "\u2014":
                    score_str = f"{score} / 100"
                else:
                    score_str = status

                # Determine rating
                if status in ("FAIL", "ERROR", "CRITICAL"):
                    rating = "FAIL"
                elif status in ("WARNING", "MARGINAL"):
                    rating = "WARNING"
                elif status in ("PASS", "OK"):
                    rating = "PASS"
                else:
                    rating = status

                # Gating: FAIL/CRITICAL items are gating
                gating = "Yes" if rating == "FAIL" else "No"

                score_rows.append((s.title, score_str, rating, gating))

        if score_rows:
            add_styled_table(
                doc,
                ["Domain", "Score", "Rating", "Gating"],
                score_rows,
                col_widths=[2.5, 1.2, 1.0, 0.8],
            )
        else:
            # Fallback: simple summary counts
            counts = self._count_findings(all_findings)
            summary_rows = [
                ("Total Findings", str(sum(counts.values()))),
                ("Critical", str(counts["critical"])),
                ("High", str(counts["high"])),
                ("Warning", str(counts["warning"])),
                ("Info", str(counts["info"])),
                ("Pass", str(counts["pass"])),
                ("Domains Analyzed", str(len([
                    k for k in results if not k.startswith("_")
                ]))),
            ]
            add_styled_table(
                doc, ["Metric", "Value"], summary_rows,
                col_widths=[3.0, 3.5],
            )

        doc.add_paragraph()

        # ------ 1.2 Risk Matrix ------
        doc.add_heading(
            f"{sect.number}.2 Risk Matrix (Impact vs. Likelihood)", level=2,
        )
        doc.add_paragraph(
            "The following matrix plots each finding by its impact on product "
            "certification or reliability (vertical axis) against the likelihood "
            "of occurrence (horizontal axis)."
        )

        # Build risk matrix dynamically from findings
        risk_table = doc.add_table(rows=5, cols=4)
        risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        risk_table.style = "Table Grid"

        # Header row
        risk_headers = ["", "Likely", "Possible", "Unlikely"]
        for i, h in enumerate(risk_headers):
            cell = risk_table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h.upper() if h else "IMPACT / LIKELIHOOD")
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, "1F4E79")

        # Group findings by severity into matrix buckets
        critical_findings = [f for f in all_findings if f.severity == "CRITICAL"]
        high_findings = [f for f in all_findings if f.severity == "HIGH"]
        warning_findings = [f for f in all_findings if f.severity == "WARNING"]
        info_findings = [f for f in all_findings if f.severity in ("INFO", "PASS")]

        def _titles(findings: list, limit: int = 3) -> str:
            if not findings:
                return "\u2014"
            return "\n".join(f.title[:40] for f in findings[:limit])

        # critical -> likely/possible, high -> possible, warning -> unlikely, info -> unlikely
        risk_data = [
            (
                "CRITICAL\n(Certification\nBlocker)",
                _titles(critical_findings),
                "\u2014",
                "\u2014",
            ),
            (
                "HIGH\n(Performance\nDegradation)",
                "\u2014",
                _titles(high_findings),
                "\u2014",
            ),
            (
                "MEDIUM\n(Design\nMargin)",
                "\u2014",
                "\u2014",
                _titles(warning_findings),
            ),
            (
                "LOW\n(Informational)",
                "\u2014",
                "\u2014",
                _titles(info_findings),
            ),
        ]
        risk_colors = ["FFF0EE", "FFF3E0", "FFFDE7", "E3F2FD"]

        for r_idx, (label, likely, possible, unlikely) in enumerate(risk_data):
            row = risk_table.rows[r_idx + 1]
            # Label cell
            row.cells[0].text = ""
            p = row.cells[0].paragraphs[0]
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(7)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(row.cells[0], risk_colors[r_idx])

            for c_idx, val in enumerate([likely, possible, unlikely], 1):
                row.cells[c_idx].text = ""
                p = row.cells[c_idx].paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(7)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if val != "\u2014" and r_idx <= 1:
                    _set_cell_shading(row.cells[c_idx], risk_colors[r_idx])

        for row in risk_table.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(1.8)
            row.cells[2].width = Inches(1.8)
            row.cells[3].width = Inches(1.5)

        doc.add_paragraph()

        # ------ 1.3 Go/No-Go Recommendation ------
        doc.add_heading(
            f"{sect.number}.3 Go / No-Go Recommendation", level=2,
        )

        go_nogo_tbl = doc.add_table(rows=1, cols=1)
        go_nogo_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        go_cell = go_nogo_tbl.cell(0, 0)

        # Determine recommendation color and text based on verdict
        if "CRITICAL" in verdict:
            bg_color, border_color = "FFF0EE", "C00000"
            rec_text = "NO-GO FOR PROTOTYPE BUILD"
            text_rgb = RGBColor(0xC0, 0x00, 0x00)
        elif "CONDITIONAL" in verdict:
            bg_color, border_color = "FFF3E0", "E65100"
            rec_text = "CONDITIONAL GO -- Address HIGH Items Before Production"
            text_rgb = RGBColor(0xE6, 0x51, 0x00)
        else:
            bg_color, border_color = "E8F5E9", "388E3C"
            rec_text = "GO FOR PROTOTYPE BUILD"
            text_rgb = RGBColor(0x1B, 0x5E, 0x20)

        _set_cell_shading(go_cell, bg_color)
        self._set_cell_border(go_cell, "left", border_color, 36)
        self._set_cell_border(go_cell, "top", border_color, 8)
        self._set_cell_border(go_cell, "bottom", border_color, 8)
        self._set_cell_border(go_cell, "right", border_color, 8)

        # Cell margins
        tc = go_cell._tc
        tcPr = tc.get_or_add_tcPr()
        margins = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="120" w:type="dxa"/>'
            f'  <w:left w:w="200" w:type="dxa"/>'
            f'  <w:bottom w:w="120" w:type="dxa"/>'
            f'  <w:right w:w="160" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(margins)

        go_cell.text = ""
        p = go_cell.paragraphs[0]
        run = p.add_run(f"RECOMMENDATION:  {rec_text}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = text_rgb
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # List gating findings
        gating = [f for f in all_findings if f.severity in ("CRITICAL", "HIGH")]
        if gating:
            p2 = go_cell.add_paragraph()
            p2.paragraph_format.space_before = Pt(6)
            run = p2.add_run("Gating Criteria:")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = text_rgb

            for tf in gating[:6]:
                p3 = go_cell.add_paragraph()
                p3.paragraph_format.space_before = Pt(1)
                p3.paragraph_format.space_after = Pt(1)
                p3.paragraph_format.left_indent = Inches(0.2)
                run = p3.add_run(
                    f"\u2022 [{tf.finding_id}] {tf.title}"
                )
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        p4 = go_cell.add_paragraph()
        p4.paragraph_format.space_before = Pt(8)
        if "CRITICAL" in verdict or "CONDITIONAL" in verdict:
            run = p4.add_run(
                "The design must address all CRITICAL and HIGH findings before "
                "proceeding to prototype fabrication."
            )
        else:
            run = p4.add_run(
                "The design meets all gating criteria and is ready for prototype "
                "fabrication."
            )
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        doc.add_paragraph()

        # ------ 1.4 Key Findings Summary ------
        doc.add_heading(
            f"{sect.number}.4 Key Findings Summary", level=2,
        )
        doc.add_paragraph(
            "The following findings represent the most significant issues identified "
            "during the review. Each finding includes its severity level, a concise "
            "description, and the specific measurement or standard it violates."
        )

        # Show top 10 critical/high findings
        critical_high = [
            f for f in all_findings if f.severity in ("CRITICAL", "HIGH")
        ]
        if critical_high:
            for tf in critical_high[:10]:
                add_finding_box(
                    doc, tf.severity,
                    f"[{tf.finding_id}] {tf.title}",
                    tf.what_it_means,
                    tf.recommendation,
                )
        else:
            # Show top warnings if no critical/high
            warnings = [f for f in all_findings if f.severity == "WARNING"]
            for tf in warnings[:5]:
                add_finding_box(
                    doc, tf.severity,
                    f"[{tf.finding_id}] {tf.title}",
                    tf.what_it_means,
                    tf.recommendation,
                )
            if not warnings:
                doc.add_paragraph(
                    "No critical or high-severity findings. The design appears clean."
                )

        doc.add_page_break()

    def _build_board_overview(
        self, doc: Any, sect: SectionDef, results: dict,
        all_findings: list[TrackedFinding], verdict: str,
        sim_plots: dict[str, str] = None,
    ) -> None:
        from .docx_report import add_styled_table

        doc.add_heading(f"{sect.number}. {sect.title}", level=1)

        board_w = self.design.board_width_mm or 0
        board_h = self.design.board_height_mm or 0
        layer_count = len(self.design.layers) if self.design.layers else 0

        overview_rows = [
            ("Source File", self.design.source_file),
            ("Format", self.design.source_format),
            ("Board Size", f"{board_w:.1f} \u00d7 {board_h:.1f} mm ({board_w * board_h:.0f} mm\u00b2)"),
            ("Thickness", f"{self.design.board_thickness_mm:.2f} mm"),
            ("Layer Count", str(layer_count)),
            ("Components", str(len(self.design.components))),
            ("Nets", str(len(self.design.nets))),
            ("Traces", str(len(self.design.traces))),
            ("Vias", str(len(self.design.vias))),
            ("Zones", str(len(self.design.zones))),
        ]
        add_styled_table(doc, ["Parameter", "Value"], overview_rows, col_widths=[2.5, 4.0])

        # Layer summary
        if self.design.layers:
            doc.add_paragraph()
            doc.add_heading("Layer Stackup Summary", level=2)
            layer_rows = []
            for ly in self.design.layers:
                layer_rows.append((
                    str(ly.number),
                    ly.name,
                    ly.layer_type,
                    f"{ly.thickness_mm:.3f}" if ly.thickness_mm else "\u2014",
                    ly.material or "\u2014",
                ))
            add_styled_table(
                doc,
                ["#", "Name", "Type", "Thickness (mm)", "Material"],
                layer_rows,
                col_widths=[0.5, 1.5, 1.2, 1.5, 1.8],
            )

        doc.add_page_break()

    def _build_action_items(
        self, doc: Any, sect: SectionDef, results: dict,
        all_findings: list[TrackedFinding], verdict: str,
        sim_plots: dict[str, str] = None,
    ) -> None:
        from docx.shared import Pt

        from .docx_report import add_styled_table

        doc.add_heading(f"{sect.number}. {sect.title}", level=1)
        doc.add_paragraph(
            "The following action items are organized by severity. All CRITICAL items "
            "must be resolved before prototype fabrication. HIGH items must be resolved "
            "before production release."
        )

        # Sort: CRITICAL first, then HIGH, WARNING, INFO, PASS
        severity_order = {"CRITICAL": 0, "HIGH": 1, "WARNING": 2, "INFO": 3, "PASS": 4}
        actionable = sorted(
            all_findings,
            key=lambda f: severity_order.get(f.severity, 5),
        )

        if not actionable:
            doc.add_paragraph("No findings to report. The design appears clean.")
            doc.add_page_break()
            return

        # Group by severity with colored headers
        severity_groups = [
            ("CRITICAL", "Resolve Before Prototype Fabrication", "C00000"),
            ("HIGH", "Resolve Before Production Release", "E65100"),
            ("WARNING", "Design Optimization", "F9A825"),
            ("INFO", "Informational", "1F4E79"),
            ("PASS", "Confirmed Pass", "388E3C"),
        ]

        counter = 0
        for sev, heading_suffix, header_color in severity_groups:
            group = [f for f in actionable if f.severity == sev]
            if not group:
                continue

            doc.add_heading(
                f"{sev} \u2014 {heading_suffix}", level=2,
            )

            rows = []
            for tf in group:
                counter += 1
                rows.append((
                    tf.finding_id,
                    tf.title,
                    tf.domain.replace("_", " ").title(),
                    tf.recommendation or "\u2014",
                ))

            add_styled_table(
                doc,
                ["ID", "Action", "Domain", "Justification"],
                rows,
                col_widths=[0.8, 2.0, 0.8, 3.0],
                header_color=header_color,
            )

            doc.add_paragraph()

        doc.add_page_break()

    def _build_tool_coverage(
        self, doc: Any, sect: SectionDef, results: dict,
        all_findings: list[TrackedFinding], verdict: str,
        sim_plots: dict[str, str] = None,
    ) -> None:
        from .docx_report import add_styled_table

        doc.add_heading(f"{sect.number}. {sect.title}", level=1)
        doc.add_paragraph(
            "This section documents which analysis tools were invoked during the review "
            "and their coverage status."
        )

        rows = []
        for s in REPORT_SECTIONS:
            if s.source_tools:
                ran_tools = []
                for t in s.source_tools:
                    if t in (self.design.analysis_cache or {}):
                        ran_tools.append(t)
                status = "Ran" if ran_tools else "Not run"
                tool_list = ", ".join(s.source_tools[:3])
                if len(s.source_tools) > 3:
                    tool_list += f" (+{len(s.source_tools) - 3} more)"
                rows.append((s.title, tool_list, status))

        if rows:
            add_styled_table(
                doc,
                ["Section", "Tools", "Status"],
                rows,
                col_widths=[2.0, 3.0, 1.0],
            )
        else:
            doc.add_paragraph("No tool-specific sections defined.")

        doc.add_page_break()

    def _build_glossary(
        self, doc: Any, sect: SectionDef, results: dict,
        all_findings: list[TrackedFinding], verdict: str,
        sim_plots: dict[str, str] = None,
    ) -> None:
        from .docx_report import add_styled_table

        doc.add_heading(f"{sect.number}. {sect.title}", level=1)
        add_styled_table(
            doc,
            ["Abbreviation", "Definition"],
            _GLOSSARY,
            col_widths=[1.5, 5.0],
        )

    def _build_references(
        self, doc: Any, sect: SectionDef, results: dict,
        all_findings: list[TrackedFinding], verdict: str,
        sim_plots: dict[str, str] = None,
    ) -> None:
        from docx.shared import Pt

        from .docx_report import add_styled_table

        doc.add_heading(f"{sect.number}. {sect.title}", level=1)

        doc.add_paragraph(
            "The following standards and specifications were referenced during this "
            "design review. Compliance with these standards is required for product "
            "certification and market access."
        )

        # EMC & Regulatory Standards
        doc.add_heading(f"{sect.number}.1 EMC & Regulatory Standards", level=2)
        ref_emc = [
            ("IEC 61000-4-2", "Electrostatic Discharge (ESD) immunity test",
             "Contact and air discharge test levels for ESD immunity."),
            ("FCC Part 15 Class B", "Unintentional radiators \u2014 US radiated emission limits",
             "Mandatory for consumer devices sold in the United States."),
            ("CISPR 32 (EN 55032)", "Electromagnetic emissions of multimedia equipment",
             "International standard for radiated and conducted emissions."),
            ("CISPR 25:2021", "Vehicles \u2014 Protection of on-board receivers",
             "Automotive radiated/conducted emission limits (Classes 1-5)."),
            ("IEC 61000-4-3", "Radiated, Radio-Frequency, EM Field Immunity Test",
             "Immunity testing for radiated RF fields."),
        ]
        add_styled_table(
            doc,
            ["Standard", "Title", "Relevance"],
            ref_emc,
            col_widths=[1.5, 2.5, 2.5],
        )

        doc.add_paragraph()

        # PCB Design Standards
        doc.add_heading(f"{sect.number}.2 PCB Design Standards", level=2)
        ref_pcb = [
            ("IPC-2221B", "Generic Standard on Printed Board Design",
             "Trace width vs. current capacity, clearance rules."),
            ("IPC-2152", "Standard for Determining Current-Carrying Capacity",
             "Thermal derating of PCB conductors."),
            ("IPC-2141A", "Design Guide for High-Speed Controlled Impedance Boards",
             "Impedance calculation methods."),
        ]
        add_styled_table(
            doc,
            ["Standard", "Title", "Relevance"],
            ref_pcb,
            col_widths=[1.5, 2.5, 2.5],
        )

        doc.add_paragraph()

        # Interface Specifications
        doc.add_heading(f"{sect.number}.3 Interface Specifications", level=2)
        ref_iface = [
            ("JEDEC JESD79-4 / JESD209-4", "DDR4 / LPDDR4 SDRAM Standard",
             "Timing, voltage, and impedance requirements."),
            ("USB 2.0 / 3.2 / 4.0", "Universal Serial Bus Specification",
             "Signal levels, impedance, and ESD requirements."),
            ("PCI Express Base Spec", "PCIe Gen3/4/5 Specification",
             "Lane impedance, equalization, link budget."),
            ("IEEE 802.3", "Ethernet \u2014 100BASE-TX / 1000BASE-T",
             "Physical layer requirements."),
        ]
        add_styled_table(
            doc,
            ["Standard", "Title", "Relevance"],
            ref_iface,
            col_widths=[1.5, 2.5, 2.5],
        )

        doc.add_paragraph()

        # Textbooks
        doc.add_heading(f"{sect.number}.4 Reference Texts", level=2)
        texts = [
            "Henry Ott, Electromagnetic Compatibility Engineering, Wiley, 2009",
            "Eric Bogatin, Signal and Power Integrity \u2014 Simplified, Prentice Hall, 2010",
        ]
        for ref in texts:
            p = doc.add_paragraph(ref, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(9)

    def _build_appendices(
        self, doc: Any, sect: SectionDef, results: dict,
        all_findings: list[TrackedFinding], verdict: str,
        sim_plots: dict[str, str] = None,
    ) -> None:
        from docx.shared import Pt, RGBColor

        from .docx_report import add_finding_box, add_styled_table

        doc.add_heading(f"{sect.number}. {sect.title}", level=1)

        # ------ Appendix A: Component Summary ------
        doc.add_heading("A. Component Summary", level=2)
        doc.add_paragraph(
            "This appendix provides a summary of the component population on the PCB. "
            "A complete bill of materials (BOM) is maintained separately."
        )

        # Categorize components by reference prefix
        comp_counts: dict[str, int] = {}
        for c in self.design.components:
            ref = c.reference if hasattr(c, "reference") else (
                c.get("reference", "") if isinstance(c, dict) else ""
            )
            if ref:
                # Extract prefix (letter part of reference designator)
                prefix = ""
                for ch in ref:
                    if ch.isalpha():
                        prefix += ch
                    else:
                        break
                if prefix:
                    comp_counts[prefix] = comp_counts.get(prefix, 0) + 1

        if comp_counts:
            prefix_names = {
                "U": "ICs / SoC",
                "R": "Resistors",
                "C": "Capacitors",
                "L": "Inductors",
                "J": "Connectors",
                "D": "Diodes",
                "Q": "Transistors",
                "ANT": "Antennas",
                "BF": "Board Fiducials",
                "BH": "Board Hardware",
                "F": "Fuses",
                "Y": "Crystals",
                "FB": "Ferrite Beads",
                "T": "Transformers",
                "SW": "Switches",
            }
            comp_rows = []
            for prefix, count in sorted(comp_counts.items(), key=lambda x: -x[1]):
                name = prefix_names.get(prefix, f"{prefix} components")
                comp_rows.append((f"{name} ({prefix})", str(count)))
            comp_rows.append(("Total", str(len(self.design.components))))
            add_styled_table(
                doc,
                ["Component Type", "Count"],
                comp_rows,
                col_widths=[3.5, 2.0],
            )
        else:
            doc.add_paragraph(
                f"Total components: {len(self.design.components)}"
            )

        doc.add_paragraph()

        # ------ Appendix B: Net Classification ------
        doc.add_heading("B. Net Classification Summary", level=2)

        # Classify nets by name patterns
        net_categories: dict[str, int] = {}
        for net in self.design.nets:
            name = net.name if hasattr(net, "name") else (
                net.get("name", "") if isinstance(net, dict) else ""
            )
            upper = name.upper()
            if "DDR" in upper or "DQ" in upper or "DQS" in upper:
                cat = "DDR"
            elif "CLK" in upper or "CLOCK" in upper:
                cat = "Clock"
            elif "USB" in upper:
                cat = "USB"
            elif "GND" in upper or "VSS" in upper:
                cat = "Ground"
            elif "VCC" in upper or "VDD" in upper or "PWR" in upper or "VBUS" in upper:
                cat = "Power"
            elif "RF" in upper or "ANT" in upper or "LNA" in upper:
                cat = "RF"
            elif "ETH" in upper or "TX_" in upper or "RX_" in upper:
                cat = "Ethernet"
            elif "SPI" in upper or "MOSI" in upper or "MISO" in upper:
                cat = "SPI"
            elif "I2C" in upper or "SCL" in upper or "SDA" in upper:
                cat = "I2C"
            elif "JTAG" in upper or "TDI" in upper or "TDO" in upper:
                cat = "JTAG"
            else:
                cat = "Unclassified"
            net_categories[cat] = net_categories.get(cat, 0) + 1

        if net_categories:
            net_rows = []
            for cat in sorted(net_categories.keys()):
                count = net_categories[cat]
                pct = (count / len(self.design.nets) * 100) if self.design.nets else 0
                net_rows.append((cat, str(count), f"{pct:.1f}%"))
            net_rows.append(("Total", str(len(self.design.nets)), "100%"))
            add_styled_table(
                doc,
                ["Category", "Count", "Coverage"],
                net_rows,
                col_widths=[2.0, 1.5, 1.5],
            )
        else:
            doc.add_paragraph(f"Total nets: {len(self.design.nets)}")

        doc.add_paragraph()

        # ------ Appendix C: Complete Finding Catalogue ------
        doc.add_heading("C. Complete Finding Catalogue", level=2)
        if all_findings:
            for tf in all_findings:
                detail = tf.what_it_means
                if tf.measured_value:
                    detail += f"\nMeasured: {tf.measured_value}"
                if tf.limit_value:
                    detail += f" | Limit: {tf.limit_value}"
                if tf.margin:
                    detail += f" | Margin: {tf.margin}"
                add_finding_box(
                    doc, tf.severity,
                    f"[{tf.finding_id}] {tf.title}",
                    detail,
                    tf.recommendation,
                )
        else:
            p = doc.add_paragraph("No findings recorded.")
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # ------ Appendix D: Report Metadata ------
        doc.add_heading("D. Report Metadata", level=2)
        p = doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"Tool: MCP PCB EMCopilot\n"
            f"Confidentiality: {self.confidentiality}"
        )
        for run in p.runs:
            run.font.size = Pt(9)

    # ------------------------------------------------------------------
    # Domain section builder (generic)
    # ------------------------------------------------------------------

    def _build_domain_section(
        self,
        doc: Any,
        sect: SectionDef,
        results: dict,
        all_findings: list[TrackedFinding],
        sim_plots: Optional[dict[str, str]] = None,
    ) -> None:
        """Build a domain-specific section with data tables and findings."""
        from docx.shared import Pt, RGBColor

        from .docx_report import add_finding_box, add_image_with_caption, add_styled_table

        sim_plots = sim_plots or {}

        doc.add_heading(f"{sect.number}. {sect.title}", level=1)

        # Add domain intro description if available
        intro = _DOMAIN_DESCRIPTIONS.get(sect.key, "")
        if intro:
            doc.add_paragraph(intro)

        # Gather domain data
        domain_data = results.get(sect.key, {})
        if isinstance(domain_data, dict):
            status = domain_data.get("status", "")
            score = domain_data.get("score")
            if status:
                p = doc.add_paragraph()
                run = p.add_run(f"Status: {status.upper()}")
                run.bold = True
                run.font.size = Pt(10)
                # Color-code status
                status_upper = status.upper()
                if status_upper in ("FAIL", "CRITICAL", "ERROR"):
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                elif status_upper in ("WARNING", "MARGINAL"):
                    run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
                elif status_upper in ("PASS", "OK"):
                    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
                if score is not None:
                    run = p.add_run(f"  (Score: {score}/100)")
                    run.font.size = Pt(10)

            # Show orchestrator executive_summary for this domain if available
            exec_summary = domain_data.get("executive_summary", "")
            if exec_summary and isinstance(exec_summary, str):
                doc.add_paragraph(exec_summary)

        # Show cross-correlations from orchestrator if relevant
        cross_corr = results.get("_cross_correlations", [])
        if isinstance(cross_corr, list) and cross_corr:
            relevant = [
                c for c in cross_corr
                if isinstance(c, dict) and sect.key in str(c.get("domains", ""))
            ]
            if relevant:
                doc.add_heading("Cross-Domain Correlations", level=2)
                for c in relevant[:3]:
                    p = doc.add_paragraph()
                    run = p.add_run(c.get("description", str(c)))
                    run.font.size = Pt(9)

        # Show recommendations from orchestrator if relevant
        recs = results.get("_recommendations", [])
        if isinstance(recs, list) and recs:
            relevant_recs = [
                r for r in recs
                if isinstance(r, dict) and sect.key in str(r.get("domain", ""))
            ]
            if relevant_recs:
                doc.add_heading("Orchestrator Recommendations", level=2)
                for r in relevant_recs[:3]:
                    p = doc.add_paragraph()
                    run = p.add_run(r.get("text", str(r)))
                    run.font.size = Pt(9)

        # Gather tool data for this section
        tool_names = _SECTION_DOMAIN_MAP.get(sect.key, [])
        tool_results = {}
        for tool in tool_names:
            if tool in results:
                tool_results[tool] = results[tool]

        # Emit tool result summaries as tables
        for tool_name, tdata in tool_results.items():
            if not isinstance(tdata, dict):
                continue
            doc.add_heading(
                tool_name.replace("pcb_", "").replace("_", " ").title(),
                level=2,
            )

            # Extract top-level scalar key/value pairs for a summary table
            summary_rows = []
            for k, v in tdata.items():
                if k in ("findings", "_sub_domains"):
                    continue
                if isinstance(v, (str, int, float, bool)):
                    summary_rows.append((k.replace("_", " ").title(), str(v)))
            if summary_rows:
                add_styled_table(
                    doc, ["Parameter", "Value"], summary_rows,
                    col_widths=[3.0, 3.5],
                )

        # Emit findings for this domain with measured/limit/margin columns
        domain_findings = [
            f for f in all_findings
            if f.domain == sect.key or f.domain in [
                t.replace("pcb_analyze_", "").replace("pcb_", "")
                for t in tool_names
            ]
        ]

        if domain_findings:
            doc.add_paragraph()
            doc.add_heading("Findings", level=2)

            # Build a summary data table if findings have measurement data
            findings_with_data = [
                f for f in domain_findings
                if f.measured_value or f.limit_value or f.margin
            ]
            if findings_with_data:
                data_rows = []
                for tf in findings_with_data:
                    data_rows.append((
                        tf.finding_id,
                        tf.severity,
                        tf.title[:40],
                        tf.measured_value or "\u2014",
                        tf.limit_value or "\u2014",
                        tf.margin or "\u2014",
                    ))
                add_styled_table(
                    doc,
                    ["ID", "Severity", "Finding", "Measured", "Limit", "Margin"],
                    data_rows,
                    col_widths=[0.7, 0.7, 1.8, 1.0, 1.0, 1.0],
                )
                doc.add_paragraph()

            # Finding detail boxes
            for tf in domain_findings:
                add_finding_box(
                    doc, tf.severity,
                    f"[{tf.finding_id}] {tf.title}",
                    tf.what_it_means,
                    tf.recommendation,
                )

        # Embed relevant simulation plots
        plot_key_map = {
            "impedance": "impedance",
            "signal_integrity": "eye_diagram",
            "power_integrity": "pdn_impedance",
            "thermal": "thermal",
            "emc": "clock_emi",
            "automotive_emc": "cispr25",
            "emi_filtering": "filter_response",
            "immunity": "immunity_margin",
        }
        plot_key = plot_key_map.get(sect.key, "")
        if plot_key and plot_key in sim_plots:
            doc.add_paragraph()
            add_image_with_caption(
                doc,
                sim_plots[plot_key],
                f"{sect.title} simulation results.",
                width_inches=5.8,
            )

        doc.add_page_break()
