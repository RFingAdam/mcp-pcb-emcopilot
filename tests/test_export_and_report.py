"""Tests for SVG export, PNG conversion, and DOCX report generation tools."""

import os
import tempfile

import pytest

from mcp_pcb_emcopilot.models.pcb_data import PCBDesignData
from mcp_pcb_emcopilot.visualization.exporter import batch_export, svg_to_file, svg_to_png

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

SIMPLE_SVG = (
    '<svg xmlns="http://www.w3.org/2000/svg" width="200" height="100">'
    '<rect x="10" y="10" width="180" height="80" fill="#4A90D9"/>'
    '<text x="100" y="55" text-anchor="middle" fill="white" font-size="14">'
    "Test PCB</text></svg>"
)


def _make_design() -> PCBDesignData:
    """Create a minimal PCBDesignData for testing."""
    d = PCBDesignData(source_file="/tmp/test_board.kicad_pcb")
    d.board_width_mm = 50.0
    d.board_height_mm = 30.0
    d.layers = [
        {"name": "F.Cu", "type": "SIGNAL"},
        {"name": "GND", "type": "SIGNAL"},
        {"name": "B.Cu", "type": "SIGNAL"},
    ]
    d.components = [
        {"reference": "U1", "value": "MCU", "x_mm": 25, "y_mm": 15, "layer": "F.Cu"},
        {"reference": "R1", "value": "10k", "x_mm": 10, "y_mm": 10, "layer": "F.Cu"},
    ]
    d.nets = [
        {"name": "GND", "index": 0},
        {"name": "VCC", "index": 1},
        {"name": "SIG1", "index": 2},
    ]
    d.traces = [
        {"net": "SIG1", "layer": "F.Cu", "width_mm": 0.15,
         "start_x": 10, "start_y": 10, "end_x": 25, "end_y": 15},
    ]
    d.vias = [
        {"x_mm": 20, "y_mm": 12, "drill_mm": 0.3, "net": "GND"},
    ]
    d.zones = []
    d.review_results = {}
    return d


# ---------------------------------------------------------------------------
# SVG exporter tests
# ---------------------------------------------------------------------------


class TestSvgExporter:
    """Tests for visualization/exporter.py functions."""

    def test_svg_to_png_creates_file(self):
        """svg_to_png should write a valid PNG file."""
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            out_path = f.name
        try:
            result = svg_to_png(SIMPLE_SVG, out_path, width=400)
            assert os.path.exists(result)
            assert os.path.getsize(result) > 100  # Non-trivial file
            # PNG magic bytes
            with open(result, "rb") as fh:
                header = fh.read(8)
            assert header[:4] == b"\x89PNG"
        finally:
            os.unlink(out_path)

    def test_svg_to_png_auto_path(self):
        """svg_to_png with no output_path should create temp file."""
        result = svg_to_png(SIMPLE_SVG)
        try:
            assert os.path.exists(result)
            assert result.endswith(".png")
        finally:
            os.unlink(result)

    def test_svg_to_file_creates_svg(self):
        """svg_to_file should write SVG content."""
        with tempfile.NamedTemporaryFile(suffix=".svg", delete=False) as f:
            out_path = f.name
        try:
            result = svg_to_file(SIMPLE_SVG, out_path)
            assert os.path.exists(result)
            with open(result) as fh:
                content = fh.read()
            assert "<svg" in content
            assert "Test PCB" in content
        finally:
            os.unlink(out_path)

    def test_batch_export_png(self):
        """batch_export should create multiple PNG files."""
        with tempfile.TemporaryDirectory() as tmpdir:
            renders = {"board": SIMPLE_SVG, "stackup": SIMPLE_SVG}
            results = batch_export(renders, tmpdir, fmt="png", width=200)
            assert len(results) == 2
            for label, path in results.items():
                assert os.path.exists(path)
                assert path.endswith(".png")

    def test_batch_export_svg(self):
        """batch_export with fmt='svg' should write SVGs."""
        with tempfile.TemporaryDirectory() as tmpdir:
            renders = {"net_gnd": SIMPLE_SVG}
            results = batch_export(renders, tmpdir, fmt="svg")
            assert "net_gnd" in results
            assert results["net_gnd"].endswith(".svg")


# ---------------------------------------------------------------------------
# DOCX report generator tests
# ---------------------------------------------------------------------------

class TestDocxReport:
    """Tests for reports/docx_report.py."""

    def test_generate_docx_basic(self):
        """generate_docx_report should produce a valid DOCX."""
        try:
            from docx import Document  # noqa: F401 — availability check only
        except ImportError:
            pytest.skip("python-docx not installed")

        from mcp_pcb_emcopilot.reports.docx_report import generate_docx_report

        design = _make_design()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            out_path = f.name
        try:
            result = generate_docx_report(
                design=design,
                session_id="test123",
                output_path=out_path,
                title="Test Report",
                subtitle="Unit Test Board",
            )
            assert os.path.exists(result)
            assert os.path.getsize(result) > 1000

            # Verify it's a valid DOCX
            doc = Document(result)
            text = "\n".join(p.text for p in doc.paragraphs)
            assert "Test Report" in text
            assert "Unit Test Board" in text
        finally:
            os.unlink(out_path)

    def test_generate_docx_with_images(self):
        """generate_docx_report should embed images from image_dir."""
        try:
            from docx import Document  # noqa: F401 — availability check only
        except ImportError:
            pytest.skip("python-docx not installed")

        from mcp_pcb_emcopilot.reports.docx_report import generate_docx_report

        design = _make_design()
        with tempfile.TemporaryDirectory() as imgdir:
            # Create a dummy PNG
            png_path = os.path.join(imgdir, "board_full.png")
            svg_to_png(SIMPLE_SVG, png_path, width=200)

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                out_path = f.name
            try:
                result = generate_docx_report(
                    design=design,
                    session_id="test456",
                    output_path=out_path,
                    image_dir=imgdir,
                )
                assert os.path.exists(result)
                # File should be larger with embedded image
                assert os.path.getsize(result) > 2000
            finally:
                os.unlink(out_path)

    def test_generate_docx_with_review_results(self):
        """generate_docx_report should include review findings."""
        try:
            from docx import Document  # noqa: F401 — availability check only
        except ImportError:
            pytest.skip("python-docx not installed")

        from mcp_pcb_emcopilot.reports.docx_report import generate_docx_report

        design = _make_design()
        design.review_results = {
            "executive_summary": {
                "overall_status": "WARNING",
                "complexity": "medium",
                "total_findings": 3,
                "total_critical": 1,
                "total_warnings": 2,
                "domains_analyzed": 2,
            },
            "domain_results": [
                {
                    "domain": "EMC",
                    "status": "WARNING",
                    "critical_count": 1,
                    "warning_count": 1,
                    "info_count": 0,
                    "findings": [
                        {
                            "severity": "CRITICAL",
                            "title": "Clock EMI failure",
                            "detail": "9th harmonic exceeds FCC",
                            "recommendation": "Add SSC",
                        }
                    ],
                },
            ],
            "emi_hotspots": [],
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            out_path = f.name
        try:
            result = generate_docx_report(
                design=design,
                session_id="test789",
                output_path=out_path,
            )
            doc = Document(result)
            text = "\n".join(p.text for p in doc.paragraphs)
            assert "WARNING" in text
            assert "EMC" in text
        finally:
            os.unlink(out_path)


# ---------------------------------------------------------------------------
# Render classification helpers
# ---------------------------------------------------------------------------

class TestRenderHelpers:
    """Tests for internal helper functions."""

    def test_classify_render_nets(self):
        """_classify_render_nets should group nets by type."""
        from mcp_pcb_emcopilot.reports.docx_report import _classify_render_nets

        design = _make_design()
        design.nets = [
            {"name": "USB_D_P"}, {"name": "USB_D_N"},
            {"name": "DDR_DQ0"}, {"name": "DDR_DQ1"},
            {"name": "RF_ANT"}, {"name": "GND"},
            {"name": "TX_P"}, {"name": "RX_N"},
            {"name": "BUCK1_1V8"},
        ]
        groups = _classify_render_nets(design)
        assert "usb" in groups
        assert "ddr" in groups
        assert "rf" in groups
        assert "ethernet" in groups
        assert "power" in groups

    def test_get_key_nets(self):
        """_get_key_nets should select representative nets."""
        from mcp_pcb_emcopilot.reports.docx_report import _get_key_nets

        design = _make_design()
        design.nets = [
            {"name": "DDR_DQ0"}, {"name": "GND"},
            {"name": "USB0_D_P"}, {"name": "RF_BLE"},
            {"name": "TX_P"}, {"name": "BUCK1_1V8"},
            {"name": "WiFi_2.4GHz"},
        ]
        keys = _get_key_nets(design)
        assert len(keys) >= 4
        assert "DDR_DQ0" in keys
        assert "GND" in keys

    def test_build_annotations_from_findings(self):
        """_build_annotations_from_findings should create annotation dicts."""
        from mcp_pcb_emcopilot.reports.docx_report import _build_annotations_from_findings

        design = _make_design()
        design.review_results = {
            "emi_hotspots": [
                {"center_x_mm": 25, "center_y_mm": 15, "risk_score": 85},
            ],
            "domain_results": [],
        }
        annots = _build_annotations_from_findings(design)
        assert len(annots) >= 1
        assert annots[0]["type"] == "warning"
        assert annots[0]["x"] == 25
